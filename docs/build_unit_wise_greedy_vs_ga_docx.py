"""Generate presentable DOCX: role of Greedy vs GA in Unit-wise scheduling."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def numbered(text: str):
        doc.add_paragraph(text, style="List Number")

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph()

    # ── Title ──────────────────────────────────────────────────────────────
    title = doc.add_heading(
        "Unit-wise Scheduler: Role of Greedy & Genetic Algorithm (GA)", 0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    r = sub.add_run(
        "Purpose, objectives, and when to use each — for peers and leadership"
    )
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run("CMF Digitalization · PPS Machine Scheduling\n").bold = True
    meta.add_run(f"Document date: {date.today().strftime('%d %B %Y')}\n")
    meta.add_run(
        "Audience: Manufacturing planners, peers, and managers. "
        "Technical jargon is kept light; a short peer appendix is included at the end.\n"
    )

    # ── 1 Executive summary ────────────────────────────────────────────────
    doc.add_heading("1. Executive summary", level=1)
    para(
        "The Unit-wise Scheduler plans work piece-by-piece (Unit 1, Unit 2, …) so that "
        "downstream operations can start on early pieces while later pieces are still "
        "upstream — matching how partial quantities move on the floor."
    )
    para(
        "Two methods can build that plan:", bold=True
    )
    bullet(
        "Greedy (Standard) — the day-to-day engine. Fast, predictable, and easy to "
        "explain. It follows part priority and places each remaining piece at the "
        "earliest sensible time on the preferred or earliest free machine."
    )
    bullet(
        "GA Research (Advanced search) — an optional improvement step. It searches many "
        "alternative sequences of unit activities under the same shop rules, then keeps "
        "a new plan only if it scores better than Greedy on agreed shop goals."
    )
    para(
        "Important for management: GA does not replace shop constraints. It tries to "
        "improve outcomes inside those constraints. If it finds no improvement, the "
        "system keeps the Greedy plan."
    )

    # ── 2 Why two methods ──────────────────────────────────────────────────
    doc.add_heading("2. Why we have two methods", level=1)
    para(
        "Scheduling many parts that share machines is a classic trade-off problem. "
        "A simple, fair rule (priority → place next piece) is trusted on the shop floor. "
        "But that same rule can leave avoidable waiting, idle gaps, or extra setups when "
        "several parts compete for the same machines."
    )
    para("So the design is intentional:")
    numbered(
        "Greedy is the reliable baseline — always available, always explainable."
    )
    numbered(
        "GA is the research / improvement layer — used when we want to hunt for a "
        "better sequence without changing business rules."
    )
    numbered(
        "Comparison (Planned | Dynamic | Unit-wise) is how we prove value on real parts."
    )

    # ── 3 Shared foundation ────────────────────────────────────────────────
    doc.add_heading("3. What both methods share (non-negotiable rules)", level=1)
    para(
        "Whether the plan comes from Greedy or GA, the same factory reality applies:"
    )
    add_table(
        ["Rule", "Plain meaning"],
        [
            [
                "Active parts only",
                "Only parts marked active for scheduling are planned",
            ],
            [
                "Route order",
                "A piece cannot start Op 20 before it finishes Op 10",
            ],
            [
                "Shifts",
                "Work is placed inside configured working shifts",
            ],
            [
                "Breakdowns / machine off",
                "Unavailable machines are not used until they are free again",
            ],
            [
                "In-progress freeze",
                "Machines already running a live job are not disrupted before now",
            ],
            [
                "Preferred machine (when set)",
                "Pinned preferred machines are respected by default",
            ],
            [
                "Partial progress",
                "Approved pieces are treated as done; only remaining pieces are re-planned",
            ],
        ],
    )
    para(
        "Batch Planned and Dynamic schedules are not overwritten. Unit-wise is a "
        "parallel plan for visibility and improvement comparison."
    )

    # ── 4 Greedy ───────────────────────────────────────────────────────────
    doc.add_heading("4. Greedy (Standard) — role, purpose, objectives", level=1)

    doc.add_heading("4.1 Purpose", level=2)
    para(
        "Give planners a fast, trustworthy unit-level plan that mirrors business "
        "priority and can be explained in one sentence: “Higher-priority parts first; "
        "each remaining piece goes to the preferred or earliest free machine as soon "
        "as the piece and machine are ready.”"
    )

    doc.add_heading("4.2 What problem it solves", level=2)
    bullet("Turn batch-style remaining work into a piece-level pipeline plan.")
    bullet("Respect part priority and operation sequence without heavy computation.")
    bullet(
        "Refresh quickly after activation, production review, or schedule rebuild."
    )
    bullet(
        "Provide a clear baseline so Advanced search has something fair to beat."
    )

    doc.add_heading("4.3 How it thinks (plain language)", level=2)
    numbered("Sort active parts by priority (then a stable tie-break).")
    numbered("Take the first part and walk its route operation by operation.")
    numbered(
        "For each remaining piece on that operation, place it at the earliest time "
        "allowed by piece readiness, machine free time, shifts, and freeze rules."
    )
    numbered(
        "Update the machine’s next free time, then continue until that part’s remaining "
        "work is fully placed."
    )
    numbered("Only then move to the next priority part.")
    para(
        "That last point is important: Greedy finishes one part’s remaining unit plan "
        "before it starts the next part. It does not “steal” idle gaps for a "
        "lower-priority part while a higher-priority part is still being laid out."
    )

    doc.add_heading("4.4 Objectives Greedy optimizes for (implicitly)", level=2)
    para(
        "Greedy does not score a multi-goal formula. Its implicit objectives are:"
    )
    add_table(
        ["Objective", "How Greedy pursues it"],
        [
            [
                "Priority fairness",
                "Higher-priority parts claim machines before lower-priority ones",
            ],
            [
                "Earliness of each placement",
                "Each piece starts as soon as piece + machine + rules allow",
            ],
            [
                "Route correctness",
                "Precedence and remaining quantity after approvals are respected",
            ],
            [
                "Speed & explainability",
                "One forward pass; planners can reason about “why this bar is here”",
            ],
            [
                "Preferred resource use",
                "Preferred / pinned machines are used when configured",
            ],
        ],
    )

    doc.add_heading("4.5 What Greedy does not try to do", level=2)
    bullet(
        "It does not globally reshuffle unit order across parts to fill every idle gap."
    )
    bullet(
        "It does not explicitly minimize setups, idle, lateness, or makespan as a "
        "weighted score."
    )
    bullet(
        "It may leave a machine idle waiting for a higher-priority piece from upstream, "
        "even if a lower-priority piece could have started sooner on that machine."
    )
    para(
        "Example peers have already seen: if Part Blue (priority first) must wait for "
        "an upstream finish before using Machine ONA, Greedy will not insert Part Red "
        "into that waiting gap unless Red has higher priority and is scheduled first."
    )

    # ── 5 GA ───────────────────────────────────────────────────────────────
    doc.add_heading("5. Genetic Algorithm / GA Research — role, purpose, objectives", level=1)

    doc.add_heading("5.1 Purpose", level=2)
    para(
        "Search for a better ordering of unit-level activities (and, where allowed, "
        "machine choices) so shared machines are used more productively — without "
        "breaking the hard shop rules listed in Section 3."
    )
    para(
        "In the product UI this appears as Advanced search / GA Research. It is "
        "optional and intended for improvement pilots, contested machine loads, and "
        "cases where Greedy looks “correct by priority” but leaves avoidable waiting "
        "or idle."
    )

    doc.add_heading("5.2 What problem it solves", level=2)
    bullet(
        "Many valid sequences exist; Greedy picks one practical sequence, not always "
        "the best-scoring one."
    )
    bullet(
        "When several parts share work centers, the order of units can change flow, "
        "setups, idle, and on-time behaviour."
    )
    bullet(
        "Management needs a controlled way to try improvements and keep them only when "
        "measurable goals improve."
    )

    doc.add_heading("5.3 How it thinks (plain language)", level=2)
    numbered(
        "Build many candidate “orderings” of the unit activities that must be done."
    )
    numbered(
        "Decode each candidate into a full unit-wise plan using the same placement "
        "rules (shifts, freeze, preferred machines, precedence)."
    )
    numbered(
        "Score each plan with a multi-objective fitness (see below)."
    )
    numbered(
        "Evolve better candidates over generations (crossover / mutation / selection)."
    )
    numbered(
        "Compare the best GA plan against the Greedy baseline; persist GA only if better."
    )

    doc.add_heading("5.4 Objectives GA is designed to improve", level=2)
    para(
        "These are the shop outcomes encoded in the research fitness. Higher fitness "
        "means a lower combined cost (with a reward for throughput)."
    )
    add_table(
        ["Objective", "Business meaning", "Direction"],
        [
            [
                "Makespan (completion window)",
                "Time from first start to last finish for the planned scope",
                "Minimize",
            ],
            [
                "Mean flow time",
                "How long pieces stay in the system on average",
                "Minimize",
            ],
            [
                "Mean waiting",
                "Gaps between successive operations for a piece",
                "Minimize",
            ],
            [
                "Mean tardiness",
                "Lateness vs due date when due dates exist",
                "Minimize",
            ],
            [
                "Setup count",
                "Number of setups implied by the sequence",
                "Minimize",
            ],
            [
                "Machine idle",
                "Empty time on machines inside the schedule window",
                "Minimize",
            ],
            [
                "Utilization gap",
                "Distance from strong effective utilization",
                "Minimize",
            ],
            [
                "Throughput",
                "Units completed per hour of the schedule window",
                "Maximize",
            ],
            [
                "Priority respect",
                "Penalty when lower-urgency work jumps ahead unfairly",
                "Minimize inversions",
            ],
        ],
    )
    para(
        "Weights for these goals can be tuned for a pilot (for example, emphasize "
        "on-time delivery more than idle). The product defaults balance flow, setups, "
        "idle, utilization, throughput, and priority."
    )

    doc.add_heading("5.5 Hard vs soft in GA", level=2)
    add_table(
        ["Type", "Examples", "Behaviour"],
        [
            [
                "Hard constraints",
                "Active scope, shifts, breakdowns, freeze, route precedence, preferred pin",
                "Never violated in a valid decoded plan",
            ],
            [
                "Soft preference",
                "Part priority ordering",
                "Encouraged via seeding and a priority-inversion penalty — not a hard block",
            ],
        ],
    )

    # ── 6 Side-by-side ─────────────────────────────────────────────────────
    doc.add_heading("6. Greedy vs GA — side-by-side", level=1)
    add_table(
        ["Aspect", "Greedy (Standard)", "GA Research (Advanced)"],
        [
            [
                "Primary role",
                "Day-to-day unit plan",
                "Optional improvement search",
            ],
            [
                "Speed",
                "Very fast (one forward pass)",
                "Slower (many candidate plans)",
            ],
            [
                "Explainability",
                "High — priority then earliest place",
                "Medium — best of many scored options",
            ],
            [
                "Part handling",
                "Part-by-part in priority order",
                "Can reshuffle unit activity order across the scope",
            ],
            [
                "Idle-gap filling across parts",
                "Generally no (by design of part-by-part pass)",
                "Possible if it improves overall score",
            ],
            [
                "Success rule",
                "Always produces the baseline plan",
                "Kept only if better than Greedy",
            ],
            [
                "Best used for",
                "Daily rebuilds, training, audits",
                "Pilots, contested machines, KPI improvement hunts",
            ],
        ],
    )

    # ── 7 How they work together ───────────────────────────────────────────
    doc.add_heading("7. How they work together in the product", level=1)
    numbered("User (or system hook) rebuilds Unit-wise.")
    numbered("Greedy always produces a complete baseline plan.")
    numbered(
        "If Advanced / GA Research is selected, GA searches and scores candidates."
    )
    numbered(
        "The better of GA vs Greedy is stored as the visible unit-wise plan "
        "(source tagged accordingly)."
    )
    numbered(
        "Planners can still compare Unit-wise against Batch Planned and Batch Dynamic "
        "on the same KPIs."
    )
    para(
        "Net effect for management: we never throw away a working Greedy plan for a "
        "weaker experiment. GA must earn its place."
    )

    # ── 8 When to use which ────────────────────────────────────────────────
    doc.add_heading("8. Guidance: when to use which", level=1)
    add_table(
        ["Situation", "Recommended method", "Why"],
        [
            [
                "Routine activation / daily refresh",
                "Greedy",
                "Fast, stable, easy to defend in standup",
            ],
            [
                "Teaching how unit-wise works",
                "Greedy",
                "Behaviour maps cleanly to priority + earliest start",
            ],
            [
                "Shared bottleneck machines with mixed parts",
                "Try GA after Greedy",
                "Sequence quality often matters more than single-part logic",
            ],
            [
                "Pilot to prove KPI gains vs Planned / Dynamic",
                "Greedy first, then GA Research",
                "Need a fair baseline, then an improvement attempt",
            ],
            [
                "Priority dispute (“why didn’t Red start in the gap?”)",
                "Start with Greedy explanation; optionally GA",
                "Greedy behaviour is intentional; GA may fill gaps if score improves",
            ],
        ],
    )

    # ── 9 Worked example ───────────────────────────────────────────────────
    doc.add_heading("9. Worked example (shared machine idle gap)", level=1)
    para(
        "Two parts share Machine ONA. Blue has higher priority and quantity 10 with a "
        "shorter cycle; Red has lower priority and quantity 4 with a longer cycle. "
        "Blue’s ONA operation waits for an upstream machine, so ONA is free early on "
        "Day 1 before Blue Unit 1 arrives."
    )
    para("Under Greedy:", bold=True)
    bullet("Blue is scheduled first because of priority.")
    bullet(
        "ONA’s early free window is not given to Red while Blue is being planned."
    )
    bullet(
        "Red starts on ONA only after Blue’s ONA queue (and machine free time) allow it."
    )
    para("If priority is swapped (Red highest):", bold=True)
    bullet("Greedy places Red first — Red Unit 1 can use that early free window.")
    bullet("Blue waits for machine availability after Red (and for its own upstream).")
    para("Under GA Research:", bold=True)
    bullet(
        "The search may consider sequences that insert other work into idle gaps if "
        "overall fitness improves (flow, idle, setups, priority penalty, etc.)."
    )
    bullet(
        "If no candidate beats Greedy on the weighted score, the Greedy plan stays."
    )

    # ── 10 Success measures ────────────────────────────────────────────────
    doc.add_heading("10. How we judge success", level=1)
    para("For a pilot part or order family, success means:")
    numbered(
        "Unit-wise Greedy already shows clearer pipeline behaviour than batch blocks "
        "where multi-operation unfinished work exists."
    )
    numbered(
        "GA Research, when used, improves one or more target KPIs versus Greedy "
        "without violating hard constraints."
    )
    numbered(
        "Side-by-side compare (Planned | Dynamic | Unit-wise) shows improvement on "
        "flow, waiting, idle, throughput, or lateness where due dates and route depth "
        "make those measures meaningful."
    )
    para(
        "Weak pilots (for example only remnant quantity on a single late operation) "
        "may not show dramatic KPI stories — choose multi-op unfinished work for fair "
        "evaluation."
    )

    # ── 11 Takeaways ───────────────────────────────────────────────────────
    doc.add_heading("11. Key takeaways for peers and managers", level=1)
    bullet(
        "Greedy = trusted daily planner. Priority-first, earliest feasible placement, "
        "easy to explain."
    )
    bullet(
        "GA = optional improver. Same shop rules; searches better sequences; kept only "
        "if better than Greedy."
    )
    bullet(
        "Idle gaps under Greedy are often a priority/sequencing consequence, not a bug."
    )
    bullet(
        "Changing priority changes Greedy behaviour immediately after rebuild; GA may "
        "further reshape sequences when worthwhile."
    )
    bullet(
        "Value is proven with the three-way compare panel, not by looking at colour "
        "bars alone."
    )

    # ── Appendix for peers ─────────────────────────────────────────────────
    doc.add_heading("Appendix A — Peer notes (light technical)", level=1)
    para(
        "This appendix is for engineering / PPS peers. Managers can skip it."
    )
    bullet(
        "Greedy implementation: sequential scope loop (priority-sorted active parts), "
        "per-operation remaining units, start = max(unit ready, machine free, now), "
        "shift placement via SchedulerEngine."
    )
    bullet(
        "GA research: activity-permutation encoding, OX crossover, swap/inversion "
        "mutation, tournament + elitism, multi-run; fitness is a weighted scalar over "
        "makespan, flow, wait, tardiness, setups, idle, util-gap, priority inversions, "
        "minus throughput reward."
    )
    bullet(
        "Persistence rule: compare GA best vs Greedy baseline; store the winner with "
        "source tag (greedy | ga)."
    )
    bullet(
        "Batch tables (planned_schedule_items / rescheduling_items) remain untouched; "
        "unit plan lives in unit_schedule_items."
    )
    bullet(
        "UI: Machine Scheduling → Unit-wise → Greedy | GA Research → Rebuild; Phase 3 "
        "compare panel for Planned | Dynamic | Unit-wise KPIs."
    )

    doc.add_heading("Appendix B — One-page decision card", level=1)
    add_table(
        ["Question", "Answer"],
        [
            [
                "What do we run every day?",
                "Greedy (Standard) rebuild",
            ],
            [
                "When do we run GA?",
                "When we want a measured improvement attempt on contested work",
            ],
            [
                "Can GA break shifts / freeze / route?",
                "No — hard constraints still apply",
            ],
            [
                "What if GA is not better?",
                "Keep Greedy; no forced downgrade",
            ],
            [
                "How do we show value to leadership?",
                "KPI compare: Planned vs Dynamic vs Unit-wise",
            ],
            [
                "Why didn’t lower-priority work fill an idle gap?",
                "Greedy schedules higher-priority parts first by design",
            ],
        ],
    )

    out = Path(__file__).resolve().parent / "UNIT_WISE_GREEDY_AND_GA_ROLES.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
