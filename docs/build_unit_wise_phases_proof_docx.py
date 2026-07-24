"""Generate UNIT_WISE_PHASES_1_TO_4_IMPLEMENTATION_PROOF.docx"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph()

    doc.add_heading("CMF Digitalization — Unit-wise Scheduling", 0)
    sub = doc.add_paragraph()
    r = sub.add_run("Implementation Proof Document: Phase 1 → Phase 4")
    r.bold = True
    r.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date.today().isoformat()}\n").italic = True
    meta.add_run("Product: PPS Machine Scheduling (Planned Schedule)\n").italic = True
    meta.add_run(
        "Scope: High-mix low-volume (HMLV) unit-wise plan alongside batch Planned & Dynamic\n"
    ).italic = True
    meta.add_run(
        "Status: Phases 1–4 implemented (pilot validation recommended next)\n"
    ).italic = True

    para(
        "This document is the single proof artifact describing what was built, "
        "where it lives in the codebase, how it works, and how to verify each phase."
    )

    # 1
    doc.add_heading("1. Executive Summary", level=1)
    para(
        "Unit-wise scheduling was delivered in four phases without replacing the existing "
        "batch Planned or Dynamic engines. Batch rescheduling_items / planned_schedule_items "
        "remain the shop’s batch view; unit_schedule_items is a parallel plan that pipelines "
        "individual units through operations."
    )
    add_table(
        ["Phase", "Name", "Status", "Core outcome"],
        [
            ["1", "MVP unit-wise greedy", "Done", "Table, rebuild API, Gantt toggle, unit bars"],
            [
                "2",
                "Shop realism",
                "Done",
                "Shifts, preferred machine pin, freeze in-progress, rework cycle-only",
            ],
            ["3", "Compare KPIs + GUI", "Done", "Planned | Dynamic | Unit-wise metrics panel"],
            [
                "4",
                "Research-grade GA",
                "Done",
                "Activity-permutation GA (OX), multi-objective fitness, multi-run stats",
            ],
        ],
    )
    para(
        "Deferred (explicitly out of phase scope): plan≡unlock look-ahead debate; "
        "NSGA-II Pareto archive; guaranteed outperformance on every remnant single-op pilot."
    )

    # 2
    doc.add_heading("2. Dual-view Architecture", level=1)
    bullet("Batch Planned — scheduling.planned_schedule_items (active ScheduleHistory)")
    bullet("Batch Dynamic — scheduling.rescheduling_items (status scheduled/rescheduled)")
    bullet("Unit-wise — scheduling.unit_schedule_items (source = greedy | ga_research)")
    para(
        "Unit-wise never mutates batch tables. Rebuild deletes/reinserts only "
        "unit_schedule_items for the active scope."
    )

    doc.add_heading("2.1 Key backend modules", level=2)
    add_table(
        ["File", "Role"],
        [
            ["backend/cmf/DB/models/scheduling.py", "UnitScheduleItem model"],
            ["backend/cmf/migrations/create_unit_schedule_items.py", "DDL / additive migration"],
            ["backend/cmf/unit_wise_scheduler.py", "Greedy simulate + rebuild entry"],
            ["backend/cmf/unit_wise_compare.py", "Phase 3 Planned/Dynamic/Unit KPIs"],
            ["backend/cmf/unit_wise_ga_research.py", "Phase 4 research GA"],
            ["backend/cmf/unit_wise_ga.py", "Thin delegate to research GA"],
            ["backend/cmf/routers/unit_wise.py", "REST API /scheduling/unit-wise/*"],
            [
                "backend/cmf/algorithm.py",
                "SchedulerEngine reused for shifts/breakdowns (batch algorithm untouched for unit-wise core)",
            ],
        ],
    )

    doc.add_heading("2.2 Key frontend modules", level=2)
    add_table(
        ["File", "Role"],
        [
            [
                "frontend/src/Pages/MachineScheduling.jsx",
                "Batch | Unit-wise mode; Greedy | GA Research; rebuild",
            ],
            ["frontend/src/components/UnitWiseComparePanel.jsx", "Phase 3 KPI GUI"],
            [
                "frontend/src/Pages/schedulingTimelineUtils.js",
                "mapUnitWiseItemsToOperations",
            ],
        ],
    )

    # 3 Phase 1
    doc.add_heading("3. Phase 1 — MVP Unit-wise (Greedy)", level=1)
    doc.add_heading("3.1 Intent", level=2)
    para(
        "Prove unit pipelining: Unit 1 can start Op20 while Unit 2 is still on Op10, "
        "stored separately from batch blocks."
    )
    doc.add_heading("3.2 Data model", level=2)
    bullet("Table: scheduling.unit_schedule_items")
    bullet(
        "Fields: order_id, part_id, unit_index, operation_id/number, machine_id, "
        "start/end, schedule_version, source, status"
    )
    bullet("Migration: create_unit_schedule_items.py (creates table or adds missing columns)")

    doc.add_heading("3.3 Greedy algorithm (what “greedy” means here)", level=2)
    para(
        "Not a classic SPT/EDD rule. It is priority list scheduling + earliest feasible start:"
    )
    bullet(
        "1. Load active IN-House parts (PartScheduleStatus = active), ordered by order-part priority"
    )
    bullet("2. For each part, walk schedulable ops in routing order")
    bullet(
        "3. Skip units already approved on that op; advance unit ready time from actual ends"
    )
    bullet(
        "4. Place remaining units one-by-one: start = max(unit_ready, machine_free, now)"
    )
    bullet("5. Update machine_free and unit_ready so downstream ops pipeline")

    doc.add_heading("3.4 APIs", level=2)
    bullet("POST /api/v1/scheduling/unit-wise/rebuild")
    bullet("GET /api/v1/scheduling/unit-wise")
    bullet("GET /api/v1/scheduling/unit-wise/parts/{part_id}")
    bullet("GET /api/v1/scheduling/unit-wise/machines/{machine_id}")
    bullet("Flag: UNIT_WISE_SCHEDULE_ENABLED (default true)")

    doc.add_heading("3.5 UI proof", level=2)
    bullet("Machine Scheduling → Planned → Segmented control Batch-wise | Unit-wise")
    bullet("Unit bars labelled U-001, U-002, …")

    # 4 Phase 2
    doc.add_heading("4. Phase 2 — Shop Realism", level=1)
    doc.add_heading("4.1 Intent", level=2)
    para(
        "Align unit-wise placement with live shop constraints used by the batch "
        "SchedulerEngine, without rewriting algorithm.py’s batch logic."
    )
    doc.add_heading("4.2 Implemented constraints", level=2)
    add_table(
        ["Constraint", "Implementation"],
        [
            [
                "Shift hours",
                "_place_within_shifts_engine via SchedulerEngine / ShiftHoursConfiguration",
            ],
            [
                "Preferred / pin machine",
                "Active job → rescheduling → operation.machine_id; UNIT_WISE_PIN_PREFERRED",
            ],
            [
                "Freeze in-progress machines",
                "_freeze_active_machines from ProductionLog inprogress",
            ],
            [
                "Rework after review",
                "First rework_due remaining slots use cycle-only (skip setup)",
            ],
            [
                "Approved units",
                "No new rows for units 1..approved; remaining re-pipelined",
            ],
        ],
    )
    doc.add_heading("4.3 Explicitly deferred", level=2)
    para(
        "Plan ≡ unlock (hard filter so next-op units only appear when upstream approved "
        "≥ unit index). Shop hard unlock on job cards remains separate."
    )

    # 5 Phase 3
    doc.add_heading("5. Phase 3 — Compare Metrics + GUI", level=1)
    doc.add_heading("5.1 Intent", level=2)
    para("Prove batch vs unit-wise value with measurable KPIs before relying on GA.")
    doc.add_heading("5.2 KPI set", level=2)
    add_table(
        ["Metric", "Purpose"],
        [
            ["Makespan", "Total completion window"],
            ["Flow / mean flow", "Time job/unit spends in system"],
            ["Waiting", "Gaps between consecutive operations"],
            ["Machine utilization", "Busy ÷ machine span (this part)"],
            ["Machine idle", "Gaps within machine span"],
            ["Throughput", "Units completed per hour"],
            ["Tardiness / Earliness", "vs order due_date when present"],
        ],
    )
    doc.add_heading("5.3 Three-way compare (final)", level=2)
    para("GET /api/v1/scheduling/unit-wise/compare?part_id=… returns:")
    bullet("batch_planned — active planned_schedule_items")
    bullet("batch_dynamic — rescheduling_items")
    bullet("unit_wise_* — latest unit_schedule_items")
    bullet("metrics_compare table with Δ vs Dynamic and Δ vs Planned")
    bullet("machines_compare per machine (e.g. Magerle util/idle)")
    doc.add_heading("5.4 GUI", level=2)
    bullet(
        "UnitWiseComparePanel on Unit-wise mode: KPI cards + Planned|Dynamic|Unit-wise "
        "table + per-machine util/idle + per-unit flows"
    )
    bullet("Pilot checklist: docs/UNIT_WISE_PHASE3_PILOT.md")

    # 6 Phase 4
    doc.add_heading("6. Phase 4 — Research-grade GA", level=1)
    doc.add_heading("6.1 Intent", level=2)
    para(
        "Optional optimizer that searches beyond myopic greedy ordering, while respecting "
        "the same hard shop constraints. GA is kept only if it beats greedy on scalar fitness."
    )
    doc.add_heading("6.2 Formulation", level=2)
    bullet("Activity a = (part, unit, operation) for each unfinished unit-op")
    bullet("Precedence: Op k before Op k+1 for the same unit")
    bullet(
        "Chromosome: permutation π of activities (+ optional machine genes when pin off)"
    )
    bullet(
        "Decode: priority-list semi-active — repeatedly schedule first precedence-feasible "
        "activity in π"
    )
    bullet(
        "Operators: tournament selection, Order Crossover (OX), swap + inversion mutation, "
        "elitism"
    )
    bullet("Multi-run with seeds → best / mean / std of Cmax + convergence series")

    doc.add_heading("6.3 Multi-objective fitness (end goals)", level=2)
    para("Cost minimized (fitness = −cost):")
    add_table(
        ["End goal", "Fitness term"],
        [
            ["Effective utilization", "Penalize util gap (100 − util)"],
            ["Minimum setups", "Penalize setup_count"],
            ["Minimum tardiness / lateness", "Penalize mean tardiness"],
            ["Avoid idle", "Penalize idle_hours_total"],
            ["Faster flow", "Penalize mean flow time"],
            ["Higher throughput", "Reward units/hour"],
            ["Respect part priority", "Soft penalty on priority inversions"],
        ],
    )
    para(
        "Note: In scheduling theory mean flow is minimized. “Maximize flow” in shop language "
        "is interpreted as faster flow + higher throughput."
    )

    doc.add_heading("6.4 Hard constraints in GA decode", level=2)
    add_table(
        ["Constraint", "How enforced"],
        [
            ["Active parts only", "Scope from PartScheduleStatus=active"],
            ["Shifts", "SchedulerEngine shift placement"],
            [
                "Breakdown / machine OFF",
                "engine._machine_next_available before place; skip permanently OFF",
            ],
            ["In-progress freeze", "_freeze_active_machines"],
            ["Routing precedence", "Activity predecessor links"],
            ["Preferred machine pin", "UNIT_WISE_PIN_PREFERRED (default true)"],
        ],
    )

    doc.add_heading("6.5 How to run", level=2)
    bullet("UI: Unit-wise → GA Research → Rebuild Unit-wise")
    bullet(
        'API: POST /scheduling/unit-wise/rebuild  { "optimizer": "ga", "part_id": <id> }'
    )
    bullet(
        "Response includes source, makespan_hours, ga.improved, ga.weights, ga.stats, "
        "ga.constraints, ga.end_goals"
    )
    bullet("Doc: docs/UNIT_WISE_PHASE4_GA.md")

    doc.add_heading("6.6 Env knobs (selected)", level=2)
    bullet(
        "UNIT_WISE_GA_POPULATION, UNIT_WISE_GA_GENERATIONS, UNIT_WISE_GA_RUNS, "
        "UNIT_WISE_GA_SEED"
    )
    bullet(
        "UNIT_WISE_GA_W_MAKESPAN, _FLOW, _WAIT, _TARD, _SETUP, _IDLE, _UTIL_GAP, "
        "_THROUGHPUT, _PRIORITY"
    )
    bullet("UNIT_WISE_PIN_PREFERRED, UNIT_WISE_GA_ENABLED")

    # 7 Completion
    doc.add_heading("7. Completion Matrix (How Much Is Done)", level=1)
    add_table(
        ["Deliverable", "Phase", "%", "Evidence"],
        [
            ["Unit schedule persistence", "1", "100%", "UnitScheduleItem + migration"],
            [
                "Greedy pipeline rebuild",
                "1–2",
                "100%",
                "simulate_unit_plan / rebuild_unit_schedule",
            ],
            ["REST APIs + feature flag", "1", "100%", "routers/unit_wise.py"],
            ["Planned Gantt unit mode", "1", "100%", "MachineScheduling.jsx toggle"],
            [
                "Shifts / pin / freeze / rework",
                "2",
                "100%",
                "unit_wise_scheduler.py Phase 2 helpers",
            ],
            [
                "KPI compare Planned|Dynamic|Unit",
                "3",
                "100%",
                "unit_wise_compare.py + UnitWiseComparePanel",
            ],
            [
                "Research GA multi-obj + OX",
                "4",
                "100%",
                "unit_wise_ga_research.py",
            ],
            [
                "GA vs greedy acceptance gate",
                "4",
                "100%",
                "keep GA only if fitness improves",
            ],
            ["Plan≡unlock alignment", "—", "0%", "Deferred"],
            ["Live multi-op pilot sign-off", "—", "Pending", "Stakeholder validation"],
        ],
    )
    para(
        "Overall phased implementation: approximately 100% of scoped Phase 1–4 engineering. "
        "Overall product readiness for default shop use: pending pilot sign-off."
    )

    # 8 Verification
    doc.add_heading("8. Verification Checklist (Proof Steps)", level=1)
    doc.add_heading("8.1 Phase 1–2", level=2)
    bullet("Set UNIT_WISE_SCHEDULE_ENABLED=true; restart API")
    bullet("POST rebuild (greedy); GET unit-wise; confirm unit_index bars on Gantt")
    bullet(
        "Confirm preferred machine stays pinned when rescheduling/active job exists"
    )
    doc.add_heading("8.2 Phase 3", level=2)
    bullet(
        "GET compare?part_id=<id>; confirm batch_planned, batch_dynamic, unit_wise, "
        "metrics_compare"
    )
    bullet("Open Unit-wise mode UI; confirm KPI panel Planned | Dynamic | Unit-wise")
    doc.add_heading("8.3 Phase 4", level=2)
    bullet(
        "Rebuild with optimizer=ga; inspect ga.improved / ga.best_objectives / "
        "ga.constraints"
    )
    bullet(
        "Prefer multi-op unfinished pilot (not single remaining op remnant) for "
        "meaningful GA vs greedy delta"
    )
    doc.add_heading("8.4 Automated tests", level=2)
    bullet("testing/test_unit_wise_compare.py")
    bullet("testing/test_unit_wise_ga_research.py (OX, fitness, setups/throughput)")
    bullet("testing/test_unit_wise_ga.py (delegate import)")

    # 9 Limitations
    doc.add_heading("9. Honest Limitations", level=1)
    bullet(
        "Unit-wise value is strongest on multi-op partial-qty HMLV flows; weak on "
        "single-op remnants."
    )
    bullet("Far due dates make tardiness/earliness uninformative.")
    bullet(
        "If pin is on and only one preferred machine, GA mainly searches unit order "
        "under multi-part contention."
    )
    bullet(
        "Outperforming Planned and Dynamic on every part is a pilot goal, not a "
        "mathematical guarantee."
    )

    # 10 Conclusion
    doc.add_heading("10. Conclusion", level=1)
    para(
        "Phases 1 through 4 of the unit-wise scheduling program are implemented in code, "
        "wired to APIs and the Machine Scheduling UI, and documented. Batch Planned and "
        "Dynamic remain intact; unit-wise is an additive dual view with greedy baseline "
        "and optional research-grade GA. The recommended next step is a controlled "
        "multi-op pilot using the Phase 3 compare panel as the acceptance record."
    )
    para("— End of implementation proof —", bold=True)

    out = Path(__file__).resolve().parent / "UNIT_WISE_PHASES_1_TO_4_IMPLEMENTATION_PROOF.docx"
    doc.save(out)
    print(str(out))


if __name__ == "__main__":
    main()
