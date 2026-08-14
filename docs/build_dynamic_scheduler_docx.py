"""Generate Dynamic Scheduler (Scheduler #3) functional & technical documentation.

Based strictly on the current repository implementation. Run from backend/cmf:

    python docs/build_dynamic_scheduler_docx.py
"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    def para(text: str, bold: bool = False, italic: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def numbered(text: str):
        doc.add_paragraph(text, style="List Number")

    def code(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        return p

    def note(text: str):
        p = doc.add_paragraph()
        r0 = p.add_run("Current implementation: ")
        r0.bold = True
        p.add_run(text)
        return p

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = str(h)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph()
        return table

    title = doc.add_heading("Dynamic Scheduler (Scheduler #3)", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    r = sub.add_run(
        "Complete functional and technical documentation of the live reconciliation scheduler as implemented in this repository"
    )
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run("CMF Digitalization · PPS Machine Scheduling\n").bold = True
    meta.add_run(f"Document date: {date.today().strftime('%d %B %Y')}\n")
    meta.add_run(
        "Scope: Describe actual code behaviour. This is not a design proposal. "
        "Where behaviour differs from earlier architecture notes, that difference is labelled “Current implementation”.\n"
    )
    meta.add_run(
        "Naming warning: class DynamicSchedulerEngine in algorithm.py is Scheduler #2 (Rescheduler), "
        "not Scheduler #3. Scheduler #3 is LiveReconciliationEngine in live_reconciliation.py.\n"
    ).bold = True

    # ------------------------------------------------------------------
    doc.add_heading("1. Big picture — three schedulers", level=1)

    para(
        "The live Gantt is not produced by a single algorithm. Three mechanisms write or seed schedule data. "
        "They share constraint helpers where possible. They do not share the same trigger or clock."
    )

    doc.add_heading("1.1 Planned Scheduler (Scheduler #1)", level=2)
    bullet("Purpose: Produce the frozen baseline plan for active in-house parts.")
    bullet("Trigger: Manual/API — POST /api/v1/scheduling/generate-schedule (routers/machine_scheduling.py → generate_schedule_endpoint).")
    bullet("Input: Active orders, part priorities, operations, machines, shifts, efficiency factor.")
    bullet("Output: scheduling.planned_schedule_items (frozen) plus an initial seed of scheduling.rescheduling_items.")
    bullet("Table: planned_schedule_items is the baseline. rescheduling_items is seeded as the day-0 live copy.")
    bullet("Automatic/manual: Manual generate. Not clock-driven.")
    bullet("Uses current time: Only as generate-time origin (datetime.now() on history). Does not watch the clock afterwards.")
    bullet("Uses production logs: Only to preserve already-completed operations when regenerating.")
    bullet("Engine: algorithm.SchedulerEngine.generate_schedule.")
    note(
        "After seeding rescheduling_items, generate_schedule also calls DynamicSchedulerEngine.dynamic_reschedule "
        "(Scheduler #2) so the live table immediately reflects remaining work. That is still Scheduler #2, not #3."
    )

    doc.add_heading("1.2 Rescheduler (Scheduler #2)", level=2)
    bullet("Purpose: Re-plan remaining quantity after production-ledger changes (approvals, remaining work, machine OFF).")
    bullet("Trigger: Production-log review, job-card complete, machine OFF/ON, admin Update, and generate-schedule Phase F.")
    bullet("Input: production_logs (approved qty, actual end), remaining/handoff qty, current rescheduling_items, machines/shifts.")
    bullet("Output: DELETE + INSERT of that part’s (or all parts’) live rows in rescheduling_items with status='rescheduled'.")
    bullet("Table: scheduling.rescheduling_items.")
    bullet("Automatic/manual: Automatic on review/complete/machine OFF; also a manual admin button.")
    bullet("Uses current time: No. Anchor is last closed log (to_date+to_time) or cascade cursor, not datetime.now().")
    bullet("Uses production logs: Yes — this is its primary input.")
    bullet("Engine: algorithm.DynamicSchedulerEngine.dynamic_reschedule. Public wrapper: algorithm.dynamic_reschedule.")
    bullet("Admin API: POST /api/v1/scheduling/dynamic-reschedule (run_dynamic_reschedule).")
    bullet("Frontend: Actual Scheduling “Update Schedule” calls this endpoint. That button is Scheduler #2, not #3.")

    doc.add_heading("1.3 Dynamic Scheduler (Scheduler #3)", level=2)
    bullet("Purpose: Reconcile the live schedule against current time, activation, and leftover past windows so Actual Scheduling does not show expired unused bars.")
    bullet("Trigger: Event hooks after activate / review / complete / machine status, plus APScheduler every 5 minutes.")
    bullet("Input: rescheduling_items (live windows) + production_logs + operation_status + datetime.now() + existing constraint helpers.")
    bullet("Output: Same table, rescheduling_items. Never planned_schedule_items. Never auto-creates a job card.")
    bullet("Automatic/manual: Automatic only. There is no Scheduler #3 API and no frontend Update button for it.")
    bullet("Uses current time: Yes. Clock is server-local naive datetime.now() (same convention as activate_job_card).")
    bullet("Uses production logs: Yes, as source of truth for actual start/end/qty and open runs.")
    bullet("Engine: live_reconciliation.LiveReconciliationEngine (subclass of DynamicSchedulerEngine).")
    bullet("Public entry: live_reconciliation.reconcile_live_schedule.")

    doc.add_heading("1.4 Why #3 exists when #2 already exists", level=2)
    para(
        "Scheduler #2 answers: “quantity changed — how much remaining work is left, and when can it run after the last log?” "
        "It does not notice that a planned live window is already in the past while nobody activated the job card. "
        "It also preserves an in-progress operation’s existing live rows, which left unused past windows frozen after a late start."
    )
    para(
        "Scheduler #3 answers: “given now and production reality, is this live bar still a feasible plan?” "
        "If the bar is entirely in the past and the job never started, it moves the live bar. "
        "If the operator just activated early or late, it rewrites remaining work from actual start. "
        "If nothing meaningful changed, it writes nothing."
    )

    add_table(
        ["", "Planned #1", "Rescheduler #2", "Dynamic #3"],
        [
            ["Clock", "Generate-time only", "Last log / cascade", "datetime.now() + actuals"],
            ["Qty remaining", "Full qty", "Yes — primary job", "Uses same remaining ledger"],
            ["Missed activation", "No", "No", "Yes — periodic"],
            ["Early/late activate", "No", "Does not run on activate", "Yes — event"],
            ["Writes planned table", "Yes", "No", "No"],
            ["Writes live table", "Seed only", "Yes", "Yes, only if needed"],
            ["Shop-floor button", "Generate", "Update Schedule", "None"],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("2. Dynamic Scheduling in simple terms", level=1)
    para(
        "For a manufacturing user: the blue/planned Gantt is the original promise. "
        "The live (Actual Scheduling) Gantt is “what we currently believe will run”. "
        "The shop does not always start or finish when the plan said. "
        "Dynamic Scheduling quietly keeps the live Gantt honest against the clock and the shop floor."
    )
    code(
        "Planned schedule (frozen)\n"
        "        +\n"
        "Production logs & job-card state (what actually happened)\n"
        "        +\n"
        "Current time on the server\n"
        "        +\n"
        "Existing machines, shifts, sequence, remaining qty\n"
        "        ↓\n"
        "Is the live Gantt still valid?\n"
        "        ↓\n"
        "If yes → do nothing\n"
        "If no  → move only the affected live bars"
    )
    para(
        "It does not punch a job card. It does not change the planned baseline. "
        "The operator still activates and completes work as today."
    )

    # ------------------------------------------------------------------
    doc.add_heading("3. When does Dynamic Scheduler come into play?", level=1)

    doc.add_heading("3.A Immediate / event-based triggers", level=2)

    para("There is no POST /live-reconcile endpoint. Events reuse existing shop-floor APIs.")

    doc.add_heading("Job-card activation", level=3)
    code(
        "POST /api/v1/scheduling/operation-status/{operation_id}/activate?operator_id=...\n"
        "  → routers/machine_scheduling.py :: activate_job_card\n"
        "  → inserts production_logs (operator_status='inprogress', from_date/from_time = datetime.now())\n"
        "  → live_reconciliation._safe_reconcile_after_event(trigger='activation', operation_id, part_id)\n"
        "  → reconcile_live_schedule(..., blocking=True)\n"
        "  → LiveReconciliationEngine.reconcile\n"
        "Rescheduler (#2) is NOT called on activate."
    )
    para(
        "What happens: activate_job_card already committed the production log (datetime.now() as from_date/from_time) "
        "before the hook runs. Scheduler #3 then puts that operation in rewrite_inprogress and activation_ops. "
        "Because the log is open, derive_operation_state is IN_PROGRESS, so remaining work is planned from actual start "
        "(early or late). The #2-style “preserve in-progress live rows” path is skipped for this op."
    )

    doc.add_heading("Production-log review", level=3)
    code(
        "PUT /api/v1/production-logs/{log_id}/status\n"
        "  → routers/production_logs.py :: update_production_log_status\n"
        "  → if status in completed / inprogress / rework / rejected:\n"
        "       algorithm.dynamic_reschedule(triggered_by_part_id=None)   # Scheduler #2, often FULL SHOP\n"
        "       _safe_reconcile_after_event(trigger='production_review', operation_id, part_id)\n"
        "  → #3 only seeds if that operation’s live max_end < now (plus same-part leftover past windows)"
    )
    para(
        "What happens: remaining qty is Scheduler #2’s job. #3 then checks leftover past windows on that op/part. "
        "If #2 already produced valid future windows, #3 no-ops."
    )

    doc.add_heading("Job-card completion", level=3)
    code(
        "POST /api/v1/scheduling/operation-status/{operation_id}/complete\n"
        "  → routers/machine_scheduling.py :: complete_job_card\n"
        "  → algorithm.dynamic_reschedule(triggered_by_part_id=None)   # Scheduler #2, FULL SHOP\n"
        "  → _safe_reconcile_after_event(trigger='completion', operation_id, part_id)\n"
        "  → #3 seeds this op and same-part ops whose live max_end < now"
    )

    doc.add_heading("Machine OFF / ON", level=3)
    code(
        "PUT /api/v1/machine-status/machine-status/{machine_id}\n"
        "  → routers/machine_status.py :: update_machine_status\n"
        "  → if status becomes OFF or previous was OFF:\n"
        "       algorithm.dynamic_reschedule(triggered_by_part_id=None)   # Scheduler #2, FULL SHOP\n"
        "       _safe_reconcile_after_event(trigger='machine_status', machine_id=...)\n"
        "  → #3 marks live ops on that machine with max_end >= now as stale_not_started seeds"
    )
    note(
        "machine_status seeding uses max_end >= now (still-live occupancy), not only past windows. "
        "That can include operations whose windows are still in the future on the affected machine. "
        "Scheduler #2 has already done a full-shop remaining-qty replan before #3 runs."
    )

    para("Events that do NOT call Scheduler #3 in the current code:")
    bullet("Out-source delivered — only Scheduler #2 (dynamic_reschedule for that part).")
    bullet("Generate schedule — Scheduler #1 then Scheduler #2 Phase F (engine.dynamic_reschedule inside the live lock), not #3.")
    bullet("Admin POST /scheduling/dynamic-reschedule — Scheduler #2 only.")
    bullet("Unit-wise rebuild — separate table (unit_schedule_items), not Scheduler #3.")
    bullet("Production-log review statuses other than completed / inprogress / rework / rejected — #2 and #3 are not called.")
    bullet("Machine status changes that are neither “now OFF” nor “previously OFF” — neither #2 nor #3.")
    para(
        "Non-periodic event hooks also add every live op on the same part_id whose max_end < now to the stale seed set "
        "(reconcile(), the part_id is not None and trigger != periodic branch). Activation of one op can therefore "
        "also move leftover expired windows on later ops of that part."
    )

    doc.add_heading("3.B Periodic / background trigger", level=2)
    code(
        "FastAPI startup (main.py :: startup_event)\n"
        "  → live_reconciliation.start_live_reconciliation_scheduler()\n"
        "  → APScheduler BackgroundScheduler\n"
        "  → interval job id='live_reconciliation'\n"
        "  → every LIVE_RECONCILE_INTERVAL_MINUTES = 5 minutes\n"
        "  → run_periodic_live_reconciliation()\n"
        "       new SessionLocal()\n"
        "       reconcile_live_schedule(trigger='periodic', blocking=False)\n"
        "       db.close()\n"
        "FastAPI shutdown → shutdown_live_reconciliation_scheduler() (wait=False)"
    )
    add_table(
        ["Item", "Actual value"],
        [
            ["Library", "APScheduler BackgroundScheduler (apscheduler>=3.10.4,<4)"],
            ["Trigger type", "interval"],
            ["Interval", "5 minutes"],
            ["Job id", "live_reconciliation"],
            ["max_instances", "1"],
            ["coalesce", "True (missed ticks collapse to one run)"],
            ["misfire_grace_time", "60 seconds"],
            ["Startup", "backend/cmf/main.py startup_event"],
            ["Disable", "LIVE_RECONCILIATION_DISABLED=1 / true / yes"],
            ["Periodic lock", "pg_try_advisory_lock — skip tick if lock held"],
            ["Overlap", "Not allowed (max_instances=1 + advisory lock)"],
            ["Restart", "Job starts again on next FastAPI startup; no queue to drain"],
        ],
    )
    para(
        "Each tick opens a fresh DB session, scans all live windows, and no-ops unless it finds a seed "
        "(past unstarted window, leftover in-progress mismatch, or leftover past remaining window)."
    )

    # ------------------------------------------------------------------
    doc.add_heading("4. Dynamic Scheduler is not activation-only", level=1)
    para(
        "Activation is one input. Documenting #3 as “activate → reschedule” is incorrect. "
        "Two independent paths exist."
    )
    add_table(
        ["Path", "When", "Typical job"],
        [
            [
                "Event-driven",
                "Immediately after activate / review / complete / machine OFF-ON",
                "React to a known shop-floor change",
            ],
            [
                "Periodic",
                "Every 5 minutes with no operator action",
                "Catch missed activation and leftover past windows if an event was missed",
            ],
        ],
    )
    para("Examples that run without an activation at that moment:")
    numbered(
        "Missed activation: live window 13 Aug 10:30–14 Aug 09:45, now 15 Aug 10:00, job card never activated. "
        "The 5-minute job sees NOT_STARTED and max_end < now, then moves the live bar."
    )
    numbered(
        "Partial remaining in the past: approved qty exists, remaining live window already ended. Periodic seeds that op."
    )
    numbered(
        "After review/complete: Scheduler #2 already ran; #3 only runs if leftover live max_end is still < now."
    )
    numbered(
        "Machine OFF/ON: #3 is invoked from machine-status, not from activate."
    )

    # ------------------------------------------------------------------
    doc.add_heading("5. Lifecycle", level=1)
    code(
        "Trigger (event hook or APScheduler)\n"
        "  → reconcile_live_schedule()                # live_reconciliation.py\n"
        "  → live_schedule_lock()                     # live_schedule_lock.py\n"
        "       blocking=True on events; try-lock on periodic\n"
        "  → LiveReconciliationEngine.reconcile()\n"
        "       now = datetime.now() unless injected\n"
        "       _live_windows()                       # rescheduling_items\n"
        "       derive_operation_state()              # production_logs + operation_status\n"
        "       classify stale / rewrite / activation seeds\n"
        "       if no seed_ops → return noop (no write)\n"
        "       _impact_part_ids()                    # parts that own seed ops\n"
        "       dynamic_reschedule(part_ids=..., live_context=...)\n"
        "            reuse SchedulerEngine helpers\n"
        "            preserve still-valid windows\n"
        "            _live_rows_equivalent() → skip DELETE/INSERT if unchanged\n"
        "  → commit only when rows actually change\n"
        "  → lock released in context manager finally"
    )

    # ------------------------------------------------------------------
    doc.add_heading("6. Operation state detection", level=1)
    para(
        "These states are derived, not stored on rescheduling_items. "
        "Function: live_reconciliation.derive_operation_state(db, operation_id, total_qty)."
    )
    para("rescheduling_items.status is only 'scheduled' | 'rescheduled' (live generation), not these states.")
    para("operation_status.status is pending / inprogress / completed (job card). Used only as a fallback.")
    para("PAUSED is not a database column. Awaiting reviewer after a submitted log is the stand-in.")

    add_table(
        [
            "Derived state",
            "How detected",
            "Move this op?",
            "Authoritative time",
            "Quantity",
            "Downstream",
        ],
        [
            [
                "COMPLETED",
                "SUM(approved_quantity) >= total_qty",
                "No",
                "MAX(to_date+to_time) via _actual_end",
                "Remaining 0",
                "Cascade from actual end",
            ],
            [
                "IN_PROGRESS",
                "Open log: operator_status=inprogress AND to_time IS NULL",
                "Not because a ticker ran. Rewrite remaining only if live start < actual start, or this is an activation event",
                "MIN(from_date+from_time) = actual start",
                "Remaining = total − approved",
                "From this op’s live/new end",
            ],
            [
                "AWAITING_REVIEW",
                "operator_status=completed, reviewer status in pending/inprogress/rework/rejected, to_time set, approved < total",
                "Only if live max_end < now (treat remaining window as stale)",
                "Closed log end; do not yank mid-review otherwise",
                "Unapproved produced qty is not treated as remaining-complete until review",
                "If remaining window moved",
            ],
            [
                "PARTIALLY_COMPLETED",
                "approved > 0 or any closed log, and not completed",
                "Only remaining; completed portion immutable. Seed if live max_end < now",
                "Last actual end for remaining start",
                "total − SUM(approved)",
                "After remaining is placed",
            ],
            [
                "NOT_STARTED",
                "No qualifying logs; operation_status not inprogress/completed-with-qty",
                "Yes if live max_end < now (missed). Early pull only on activation event",
                "now (missed) or actual activation time (early/late activate)",
                "Full remaining / handoff qty",
                "Later ops on the same part if this op’s end changes and they are no longer feasible",
            ],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("7. Scheduling decision logic", level=1)

    doc.add_heading("Q1 — When is the schedule stale?", level=2)
    bullet("NOT_STARTED and live window max_end < now → missed activation seed.")
    bullet("IN_PROGRESS and live min_start < actual_start from logs → unused past window after late start.")
    bullet("PARTIALLY_COMPLETED or AWAITING_REVIEW and live max_end < now → remaining window expired.")
    bullet("Activation event: the activated operation is always a rewrite seed (early or late).")
    bullet("machine_status event: live ops on that machine with max_end >= now are seeded (see note in §3).")

    doc.add_heading("Q2 — When is it a no-op?", level=2)
    numbered("No seed_ops after the scan → return immediately; no DELETE/INSERT; no version bump.")
    numbered("Seeds exist, engine proposes rows, _live_rows_equivalent is True for every impacted part → skip writes; result.noop True if nothing written.")
    numbered("Periodic lock not acquired → skipped=True, no work.")
    para(
        "Equivalence key (per part): (operation_id, machine_id, start/end truncated to seconds, total_qty, completed_qty, remaining_qty). "
        "id and schedule_version are ignored."
    )

    doc.add_heading("Q3 — Which operations are affected?", level=2)
    para(
        "Seeds → parts that own those operations (LiveReconciliationEngine._impact_part_ids). "
        "The engine then walks every operation on those parts in sequence."
    )
    note(
        "Earlier design mentioned “colliding machines → neighboring parts”. "
        "The current _impact_part_ids returns only seed parts. Other parts are not rewritten. "
        "Their future occupancy is applied as a machine pre-block via _preblock_out_of_scope_live_rows "
        "so in-scope work does not steal those slots."
    )

    doc.add_heading("Q4 — New earliest feasible start?", level=2)
    bullet("Missed / stale pending: max(cascade_cursor, now), then _select_machine → adjust_to_shift / machine OFF windows.")
    bullet("Activation: max(cascade_cursor, actual_start from open log). Early start cannot beat predecessor/machine/shift.")
    bullet("Late activation rewrite: remaining from actual_start (not from the old past bar, not slid to ticker now).")
    bullet("Still-valid pending (end >= now and start >= cascade): keep existing live start — do not pull to now.")

    doc.add_heading("Q5 — Downstream?", level=2)
    para(
        "On the same part, later operations are visited in sequence. "
        "If still valid (end >= now and start >= new cascade), they are preserved. "
        "If the predecessor’s new end pushes into their window (start < cascade), they are replanned from cascade. "
        "Other parts are not in the impact set."
    )

    # ------------------------------------------------------------------
    doc.add_heading("8. How “now” is used", level=1)
    add_table(
        ["Timestamp", "Source", "Role"],
        [
            ["Current time", "datetime.now() naive local (injectable in tests)", "Missed-activation floor; stale remaining floor"],
            ["Activation / actual start", "production_logs.from_date + from_time", "IN_PROGRESS and activation rewrite anchor"],
            ["Actual end", "MAX(to_date+to_time)", "Completed / closed-log cascade cursor (_actual_end)"],
            ["Planned start/end", "planned_schedule_items", "Baseline only — never written by #3"],
            ["Live start/end", "rescheduling_items", "What #3 validates and may replace"],
        ],
    )
    bullet("NOT_STARTED missed: now can become the scheduling floor.")
    bullet("NOT_STARTED still in the future: now is ignored; live window kept.")
    bullet("IN_PROGRESS: actual start is authoritative; ticker must not slide the start to now.")
    bullet("COMPLETED: actual completion is authoritative; #3 does not move that op.")

    # ------------------------------------------------------------------
    doc.add_heading("9. Early activation", level=1)
    para("Example: planned/live start 15/08 14:00, operator activates 15/08 12:30.")
    numbered("Event: POST .../operation-status/{id}/activate → activate_job_card writes from_date/from_time = 12:30.")
    numbered("Hook: _safe_reconcile_after_event(trigger='activation').")
    numbered("Detection: trigger==activation adds operation_id to activation_ops and rewrite_inprogress.")
    numbered("Anchor: actual start 12:30, then max(predecessor cascade, 12:30), then machine/shift.")
    numbered("Move earlier?: Yes if constraints allow. Periodic ticks never pull unstarted work earlier.")
    numbered("Constraints: same _select_machine, adjust_to_shift, handoff qty, pin/alternative tiers.")
    numbered("Downstream: later ops on the same part replan only if their live start is now before the new cascade.")
    numbered("Write: DELETE+INSERT that part’s scheduled/rescheduled rows if not equivalent; bump schedule_version on new rows.")

    # ------------------------------------------------------------------
    doc.add_heading("10. Late activation", level=1)
    code(
        "Operator activates\n"
        "  → production_logs actual start = 15/08 10:00\n"
        "  → Scheduler #3 (not #2)\n"
        "  → live min_start (13/08) < actual start → rewrite_inprogress\n"
        "  → skip #2-style preserve of the old 13–14 Aug bar\n"
        "  → remaining work from actual start via _schedule_remaining_work_blocks\n"
        "  → downstream on same part if no longer feasible\n"
        "  → rescheduling_items DELETE+INSERT for that part\n"
        "  → Gantt reads the new live rows"
    )
    para(
        "The old past live window is not frozen merely because the op became IN_PROGRESS. "
        "That preserve path is skipped when the op is in rewrite_inprogress_op_ids. "
        "After the rewrite, later periodic ticks see live start >= actual start and do not move the started op."
    )

    # ------------------------------------------------------------------
    doc.add_heading("11. Missed activation", level=1)
    para("Example: live 13/08 10:30–14/08 09:45, now 15/08 10:00, job card never activated.")
    bullet("Detected: yes. Periodic scan: NOT_STARTED and max_end < now.")
    bullet("Trigger: APScheduler (not activation). An event is not required.")
    bullet("Moves live schedule: yes, earliest feasible start >= now, same constraint engine.")
    bullet("Auto-activate job card: no. No production_logs row is inserted.")
    bullet("Planned baseline: unchanged.")
    bullet("Job-card status: stays pending / absent. Operator must still activate to work.")
    para(
        "This is schedule validity, not shop-floor permission. "
        "activate_job_card still enforces leave, pending review, upstream qty, and breakdown blocks."
    )

    # ------------------------------------------------------------------
    doc.add_heading("12. Early completion", level=1)
    para("Example: planned 10:00–14:00, actual 10:00–12:00.")
    bullet("Who handles it: job-card complete and/or production-log review → Scheduler #2 first (often full shop).")
    bullet("Completed op: skipped for insertion; cascade_cursor = actual end (12:00).")
    bullet("Downstream: #2 may pull later ops earlier from that cursor, subject to constraints.")
    bullet("Scheduler #3 after that: only if leftover live max_end is still < now; otherwise no-op.")
    bullet("Authoritative timestamp: production log to_date+to_time, not a scheduler-invented finish.")

    # ------------------------------------------------------------------
    doc.add_heading("13. Late completion", level=1)
    para("Example: planned 10:00–14:00, actual 10:00–16:30.")
    bullet("Scheduler #2 runs on complete/review and pushes downstream from actual end 16:30 through shift/machine/sequence helpers.")
    bullet("Scheduler #3 runs afterwards for leftover past windows on that part.")
    bullet("Duplicate full rebuild: #2 may already have rewritten the shop (triggered_by_part_id=None). #3 then typically no-ops via empty seeds or row equivalence.")
    bullet("Lock: both share LIVE_SCHEDULE_LOCK_KEY; same session is re-entrant so they do not deadlock.")

    # ------------------------------------------------------------------
    doc.add_heading("14. Partial production", level=1)
    code(
        "total_qty = 5\n"
        "SUM(approved_quantity) = 4   → completed portion immutable\n"
        "remaining = 1                → _schedule_remaining_work_blocks / remaining-work path\n"
        "Multiple logs: approved is summed across all logs for the operation."
    )
    para(
        "Handoff to the next operation uses production_log_helpers.get_operation_handoff / get_operation_work_due "
        "(upstream approved, rework due, reject due). Shop-floor activate is still hard-capped by upstream approved qty."
    )
    para(
        "If remaining live window max_end is still in the future, #3 preserves it. "
        "If that remaining window is already in the past, #3 seeds it and places remaining qty from max(actual_end, now)."
    )

    # ------------------------------------------------------------------
    doc.add_heading("15. Production logs interaction — source of truth", level=1)
    add_table(
        ["Store", "Table", "Represents"],
        [
            ["Actual production reality", "scheduling.production_logs", "Actual start/end, produced/approved/rework/reject qty, open run"],
            ["Job-card status", "scheduling.operation_status", "pending / inprogress / completed timestamps (fallback for derive_operation_state)"],
            ["Live schedule", "scheduling.rescheduling_items", "Current Gantt bars — a plan, not proof of execution"],
            ["Frozen baseline", "scheduling.planned_schedule_items", "Original generate-schedule promise"],
        ],
    )
    para(
        "Scheduler #3 must never overwrite actual production timestamps with generated ones. "
        "It only writes live plan rows."
    )

    # ------------------------------------------------------------------
    doc.add_heading("16. rescheduling_items lifecycle", level=1)
    code(
        "generate-schedule\n"
        "  → wipe/seed status='scheduled' (completed ops preserved)\n"
        "  → Scheduler #2 often converts rows to status='rescheduled'\n"
        "Rescheduler #2 or Dynamic #3 (per impacted part)\n"
        "  → if live_context and rows equivalent → skip\n"
        "  → else DELETE part’s scheduled|rescheduled rows\n"
        "  → INSERT new rows status='rescheduled', schedule_version = MAX+1\n"
        "Gantt reads current rows only"
    )
    bullet("Not a history table. Old versions are deleted. Latest generation is the only live state.")
    bullet("schedule_version increments only when new rows are inserted.")
    bullet("No write: no version bump, ids stay the same.")
    bullet("It is the current live state for Actual Scheduling.")

    # ------------------------------------------------------------------
    doc.add_heading("17. Planned vs live", level=1)
    code(
        "PLANNED  (planned_schedule_items, GET /scheduling/gantt-data)\n"
        "  13/08 10:30 → 14/08 09:45     ← never moved by #3\n"
        "\n"
        "LIVE     (rescheduling_items, GET /scheduling/view-rescheduling)\n"
        "  after missed activation at 15/08 10:00:\n"
        "  15/08 10:00 → calculated end  ← only this moves"
    )
    para(
        "Planned remains the baseline for comparison. Live is the working timetable. "
        "That is why Actual Scheduling can disagree with Planned Scheduling without regenerating the baseline."
    )

    # ------------------------------------------------------------------
    doc.add_heading("18. Constraint handling", level=1)
    para("Scheduler #3 does not implement a third algorithm. It passes live_context into DynamicSchedulerEngine.dynamic_reschedule.")

    add_table(
        ["Constraint", "Where", "Reused by #3?"],
        [
            ["Pinned machine / WC alternative / best WC", "DynamicSchedulerEngine._select_machine", "Yes"],
            ["Machine OFF windows", "SchedulerEngine._machine_next_available (MachineStatus status_id=2)", "Yes"],
            ["Machine conflicts (in-scope)", "self.machine_end_time occupancy", "Yes"],
            ["Out-of-scope live occupancy", "_preblock_out_of_scope_live_rows", "Yes (live mode only)"],
            ["Shift window / working day / OT via assignments", "adjust_to_shift, _shift_window, MachineOperatorShiftAssignment", "Yes"],
            ["Operation sequence", "ops ordered by operation_number; cascade_cursor", "Yes"],
            ["Handoff / released qty", "production_log_helpers.get_operation_handoff", "Yes"],
            ["Remaining / rework / reject duration", "_schedule_remaining_work_blocks, get_operation_work_due", "Yes"],
            ["Setup vs cycle", "_operation_duration_hours; setup_already_done flag", "Yes"],
            ["Completed qty immutable", "_is_operation_completed / approved sum", "Yes"],
            ["Operator leave / pending-review activate blocks", "activate_job_card only", "Not inside #3 — #3 never auto-activates"],
        ],
    )
    note(
        "Operator Gantt occupancy is not a separate Scheduler #3 constraint beyond existing shift assignment "
        "lookups already used by adjust_to_shift. #3 does not invent new operator conflict logic."
    )

    # ------------------------------------------------------------------
    doc.add_heading("19. Impact set", level=1)
    para("Current implementation does not reschedule the entire shop on a Dynamic tick.")
    code(
        "seed operations (stale / rewrite / activation)\n"
        "  → parts that own those operations\n"
        "  → walk remaining operations on those parts only\n"
        "  → preserve still-valid ops on that part\n"
        "  → replan only stale / blocked-by-predecessor ops\n"
        "Other parts: not rewritten; their future machine use is pre-blocked"
    )
    para(
        "Example: Part A op 10 window is entirely in the past (not started). Part B on another machine is still in September. "
        "Only Part A is in impact_parts. Part B keeps the same row ids."
    )
    note(
        "The live_reconciliation.py module docstring still mentions “colliding machines”. "
        "_impact_part_ids currently returns seed parts only. Collision handling is pre-block, not rewrite."
    )

    # ------------------------------------------------------------------
    doc.add_heading("20. Idempotency", level=1)
    code(
        "10:00  live generated / last good reconcile\n"
        "10:05  periodic: no seed → NO-OP\n"
        "10:10  periodic: no seed → NO-OP\n"
        "10:15  production event or missed window → reconcile, maybe write\n"
        "10:20  windows now valid → NO-OP"
    )
    para("Mechanisms:")
    bullet("Empty seed set short-circuits before planning.")
    bullet("Missed activation uses max_end < now, not start < now, so a mid-window unstarted bar is not rewritten every tick.")
    bullet("After a move, the new end is usually still in the future, so the next tick has no seed.")
    bullet("_live_rows_equivalent skips DELETE/INSERT when proposed equals current.")
    bullet("IN_PROGRESS after a successful late-start rewrite is not moved by the ticker.")

    # ------------------------------------------------------------------
    doc.add_heading("21. Locking and concurrency", level=1)
    bullet("Lock: PostgreSQL session-level advisory lock, key 0x4C495645 (“LIVE”).")
    bullet("Module: live_schedule_lock.py — live_schedule_lock / pg_advisory_lock / pg_try_advisory_lock.")
    bullet(
        "Holders: public algorithm.dynamic_reschedule wrapper (#2), reconcile_live_schedule (#3), "
        "and generate-schedule seed + Phase F. Phase F calls DynamicSchedulerEngine.dynamic_reschedule "
        "directly (engine method, no second lock). LiveReconciliationEngine.reconcile also calls the engine method "
        "while already holding the lock from reconcile_live_schedule."
    )
    bullet("Scope: one live-table writer at a time across backend instances sharing the database.")
    bullet("Periodic: try-lock; if busy, skip the tick.")
    bullet("Events: blocking lock. Same DB session may re-enter (review runs #2 then #3).")
    bullet("SQLite tests: lock is a no-op that always acquires.")
    para("uvicorn --reload still needs this lock (reloader + worker, or multiple API processes).")

    # ------------------------------------------------------------------
    doc.add_heading("22. Failure recovery", level=1)
    bullet("Event hook: _safe_reconcile_after_event swallows exceptions so activate/complete/review still succeed.")
    bullet("reconcile_live_schedule crash: rollback, return success=False; originating action already committed.")
    bullet("Periodic failure: logged; next 5-minute tick retries (idempotent).")
    bullet("Backend / APScheduler restart: job is registered again on startup; no persistent job queue.")
    bullet("Partial writes: dynamic_reschedule wraps in try/except with rollback on failure.")
    bullet("Live empty scope: #3 must not wipe the whole rescheduling_items table (guard when live_context is set).")

    # ------------------------------------------------------------------
    doc.add_heading("23. Frontend / Gantt flow", level=1)
    code(
        "Scheduler #3 writes rescheduling_items\n"
        "  → GET /api/v1/scheduling/view-rescheduling\n"
        "       routers/machine_scheduling.py :: view rescheduling (gantt payload)\n"
        "  → frontend/src/Pages/ActualScheduling.jsx :: fetchSchedule()\n"
        "  → SchedulingGanttTimeline"
    )
    bullet("No websocket / no event push from #3.")
    bullet("No polling loop. fetchSchedule runs on page load (useEffect + setTimeout 0).")
    bullet("Refresh after Scheduler #2 Update Schedule only (handleUpdateSchedule).")
    bullet("No last-updated timestamp from #3. schedule_version is on rows but the Gantt fetch does not highlight version changes.")
    bullet("Planned page uses GET /scheduling/gantt-data (planned_schedule_items) and is unaffected.")
    note(
        "The Actual Scheduling “Update Schedule” button is Scheduler #2. "
        "There is no Dynamic Scheduler button, by product decision."
    )

    # ------------------------------------------------------------------
    doc.add_heading("24. End-to-end examples", level=1)

    examples = [
        (
            "Example 1 — On-time execution",
            "Activate on time; produce; complete.",
            "Activation event then later complete/review.",
            "IN_PROGRESS then COMPLETED.",
            "If live window already matched actual start, equivalence / preserve → little or no #3 write on activate. Complete → #2 cascades; #3 no-ops if windows valid.",
            "Actual start; then actual end.",
            "Shift/machine/sequence as usual.",
            "Complete: often full shop via #2. #3 only leftover past windows on that part.",
            "Live rows follow remaining work. Planned unchanged.",
            "Refresh Actual Scheduling to see it.",
        ),
        (
            "Example 2 — Late activation",
            "Window 13–14 Aug unused; activate 15 Aug 10:00.",
            "POST activate → #3 trigger=activation.",
            "IN_PROGRESS; live min_start < actual start.",
            "Do not preserve old bar. Remaining from 15 Aug 10:00.",
            "actual_start.",
            "_select_machine, shifts, predecessor.",
            "That part; downstream ops on the same part if start < new cascade.",
            "Old 13–14 Aug rows deleted; new remaining rows inserted.",
            "Gantt shows remaining from 15 Aug.",
        ),
        (
            "Example 3 — Missed activation",
            "Same window expired; never activated; 15 Aug 10:00.",
            "Periodic APScheduler.",
            "NOT_STARTED, max_end < now.",
            "Move live bar to feasible >= now. Do not insert a production log.",
            "now.",
            "Same engine constraints.",
            "That part only (plus same-part ops that become infeasible).",
            "Live rows replaced for that part. Planned untouched. Job card still pending.",
            "Next Gantt refresh shows the moved bar.",
        ),
        (
            "Example 4 — Early activation",
            "Live 15/08 14:00; activate 12:30.",
            "POST activate → #3.",
            "IN_PROGRESS rewrite via activation_ops.",
            "Start at 12:30 if cascade/machine/shift allow; else later.",
            "actual_start, then constraints.",
            "Predecessor and machine occupancy can block true early start.",
            "Same part.",
            "New live start ≤ original if feasible.",
            "Gantt may show an earlier bar.",
        ),
        (
            "Example 5 — Late completion",
            "Finish 16:30 instead of 14:00.",
            "Complete and/or log review → #2 then #3.",
            "COMPLETED; downstream pending.",
            "#2 places downstream from 16:30. #3 no-ops if already valid.",
            "actual end.",
            "Shifts, machines, sequence.",
            "#2 may be full shop; #3 scoped to leftover past windows.",
            "Downstream live rows later.",
            "Gantt shows delay propagated.",
        ),
        (
            "Example 6 — Partial production",
            "Qty 5; 4 approved; 1 remaining.",
            "Review → #2 remaining-work; #3 if remaining window already ended.",
            "PARTIALLY_COMPLETED.",
            "Completed 4 frozen; 1 remaining scheduled.",
            "Last actual end, or now if that remaining window is past.",
            "setup_already_done typically true after logs exist.",
            "That part’s remaining chain.",
            "Live remaining_qty reflects leftover units.",
            "Gantt remaining bar only.",
        ),
    ]
    for title_e, situation, trigger, state, decision, anchor, constraints, impact, dbw, gantt in examples:
        doc.add_heading(title_e, level=2)
        add_table(
            ["Step", "What happens"],
            [
                ["Situation", situation],
                ["Trigger", trigger],
                ["State", state],
                ["Decision", decision],
                ["Anchor", anchor],
                ["Constraints", constraints],
                ["Impact", impact],
                ["Database", dbw],
                ["Gantt", gantt],
            ],
        )

    # ------------------------------------------------------------------
    doc.add_heading("25. Scheduler decision matrix", level=1)
    add_table(
        [
            "Situation",
            "Trigger",
            "Scheduler",
            "Action",
            "Anchor",
            "Moves current op?",
            "Moves downstream?",
        ],
        [
            [
                "On-time activation, live already matches",
                "Activate",
                "#3",
                "Rewrite flagged; often equivalent / preserve",
                "Actual start",
                "Only if live bar disagrees",
                "Only if this op’s end changes",
            ],
            [
                "Early activation",
                "Activate",
                "#3",
                "Reconcile remaining from actual start",
                "Actual start",
                "Yes if constraints allow earlier",
                "Same part if cascade requires",
            ],
            [
                "Late activation",
                "Activate",
                "#3",
                "Drop unused past bar; remaining from actual start",
                "Actual start",
                "Yes (remaining window)",
                "Same part if needed",
            ],
            [
                "Missed activation",
                "Periodic 5 min",
                "#3",
                "Move live bar; do not activate",
                "now",
                "Yes (live only)",
                "Same part if no longer feasible",
            ],
            [
                "Early completion",
                "Complete / review",
                "#2 then #3",
                "#2 cascade; #3 leftover past only",
                "Actual end",
                "No (completed frozen)",
                "Yes via #2",
            ],
            [
                "Late completion",
                "Complete / review",
                "#2 then #3",
                "Same",
                "Actual end",
                "No",
                "Yes via #2",
            ],
            [
                "Partial production",
                "Log review",
                "#2 then #3",
                "Remaining qty via #2; #3 if remaining window past",
                "Actual end / now if stale",
                "Remaining only",
                "Same part if remaining end changes",
            ],
            [
                "Machine OFF/ON",
                "Machine status",
                "#2 (full shop) then #3",
                "#2 replans remaining; #3 seeds ops on that machine with max_end >= now",
                "#2: logs/cascade; #3: live+now",
                "Possibly via #2",
                "Often yes via #2",
            ],
            [
                "No state change, windows still valid",
                "Periodic",
                "#3",
                "No-op",
                "—",
                "No",
                "No",
            ],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("26. File / function reference", level=1)
    add_table(
        ["Responsibility", "File", "Function / class"],
        [
            ["Dynamic Scheduler (#3)", "backend/cmf/live_reconciliation.py", "LiveReconciliationEngine, reconcile_live_schedule"],
            ["APScheduler startup/shutdown", "backend/cmf/main.py", "startup_event / shutdown_event"],
            ["APScheduler job", "backend/cmf/live_reconciliation.py", "start_live_reconciliation_scheduler, run_periodic_live_reconciliation"],
            ["Activation trigger", "backend/cmf/routers/machine_scheduling.py", "activate_job_card → _safe_reconcile_after_event"],
            ["Completion trigger", "backend/cmf/routers/machine_scheduling.py", "complete_job_card"],
            ["Production-log trigger", "backend/cmf/routers/production_logs.py", "update_production_log_status"],
            ["Machine trigger", "backend/cmf/routers/machine_status.py", "update_machine_status"],
            ["State detection", "backend/cmf/live_reconciliation.py", "derive_operation_state"],
            ["Live windows", "backend/cmf/live_reconciliation.py", "_live_windows"],
            ["Impact calculation", "backend/cmf/live_reconciliation.py", "LiveReconciliationEngine._impact_part_ids"],
            ["Scheduling engine (#2 reused)", "backend/cmf/algorithm.py", "DynamicSchedulerEngine.dynamic_reschedule"],
            ["Live context flag", "backend/cmf/algorithm.py", "LiveReconcileContext"],
            ["Machine selection", "backend/cmf/algorithm.py", "_select_machine, _pick_best_machine, _machine_next_available"],
            ["Shift calculation", "backend/cmf/algorithm.py", "adjust_to_shift, _shift_window"],
            ["Remaining work", "backend/cmf/algorithm.py", "_schedule_remaining_work_blocks, _schedule_operation_blocks"],
            ["Equivalence / no-op write", "backend/cmf/algorithm.py", "_live_rows_equivalent"],
            ["Out-of-scope occupancy", "backend/cmf/algorithm.py", "_preblock_out_of_scope_live_rows"],
            ["Handoff / remaining qty", "backend/cmf/production_log_helpers.py", "get_operation_handoff, get_operation_work_due, total_approved_for_operation"],
            ["Advisory lock", "backend/cmf/live_schedule_lock.py", "live_schedule_lock, LIVE_SCHEDULE_LOCK_KEY"],
            ["Rescheduler public API (#2)", "backend/cmf/algorithm.py + routers/machine_scheduling.py", "dynamic_reschedule / run_dynamic_reschedule (POST /dynamic-reschedule)"],
            ["Planned generate (#1)", "backend/cmf/algorithm.py + routers/machine_scheduling.py", "SchedulerEngine.generate_schedule / generate_schedule_endpoint"],
            ["Live Gantt API", "backend/cmf/routers/machine_scheduling.py", "view_rescheduling_items (GET /view-rescheduling)"],
            ["Planned Gantt API", "backend/cmf/routers/machine_scheduling.py", "GET /gantt-data"],
            ["Actual Scheduling UI", "frontend/src/Pages/ActualScheduling.jsx", "fetchSchedule, handleUpdateSchedule (#2 only)"],
            ["Live table model", "backend/cmf/DB/models/scheduling.py", "Rescheduling, ProductionLog, OperationStatus"],
        ],
    )

    # ------------------------------------------------------------------
    doc.add_heading("27. When does Scheduler #3 come into play? — final answer", level=1)
    para("Remember it this way:", bold=True)
    para(
        "Scheduler #3 comes into play automatically whenever the live Gantt may be lying about time or activation — "
        "not whenever quantity changed, and not when the user clicks Update Schedule."
    )
    doc.add_heading("EVENT-DRIVEN", level=2)
    para(
        "Immediately after job-card activate, production-log review, job-card complete, or machine OFF/ON. "
        "Activation is the path that handles early/late start. Review/complete mainly let #2 do remaining-qty work; "
        "#3 only cleans leftover past live windows."
    )
    doc.add_heading("PERIODIC", level=2)
    para(
        "Every 5 minutes the API process scans live bars. If an unstarted (or leftover remaining) window has already ended, "
        "it moves that live bar to a feasible time now. It never punches the job card. If every live bar is still valid, it writes nothing."
    )
    para("It does not come into play for:", bold=True)
    bullet("Generating the planned baseline (#1).")
    bullet("The Actual Scheduling Update button (#2).")
    bullet("A clock tick with no stale windows and no events (no-op).")

    # ------------------------------------------------------------------
    doc.add_heading("28. Implementation notes vs earlier design", level=1)
    numbered(
        "No Scheduler #3 HTTP API exists. Events piggy-back on existing endpoints."
    )
    numbered(
        "Impact set is seed parts only; colliding parts are pre-blocked, not rewritten."
    )
    numbered(
        "Still-valid live windows (end >= now) are preserved and not pulled to now."
    )
    numbered(
        "Missed activation uses entire window in the past (max_end < now), not merely start < now."
    )
    numbered(
        "Review/complete still call Scheduler #2 with triggered_by_part_id=None (full shop) before #3."
    )
    numbered(
        "Frontend does not poll; users refresh or reopen Actual Scheduling to see #3 results."
    )
    numbered(
        "Clock is naive datetime.now() on the API server, matching activate_job_card."
    )
    numbered(
        "unit_schedule_items / unit-wise scheduler is a separate view and is not updated by #3."
    )
    numbered(
        "Activation commits the production log first; #3 cannot roll back the job-card activation if reconciliation fails."
    )
    numbered(
        "Public algorithm.dynamic_reschedule acquires the advisory lock. Engine.dynamic_reschedule does not; "
        "callers that already hold the lock (Phase F, LiveReconciliationEngine) use the engine method."
    )

    para(
        "This document describes the code as of the document date. "
        "If live_reconciliation.py or the event hooks change, regenerate this file from docs/build_dynamic_scheduler_docx.py."
    )

    out = Path(__file__).resolve().parent / "DYNAMIC_SCHEDULER_HOW_IT_WORKS.docx"
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
