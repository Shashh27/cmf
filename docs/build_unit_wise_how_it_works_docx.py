"""Generate manager-friendly Unit-wise how-it-works DOCX (non-technical)."""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def numbered(text: str):
        doc.add_paragraph(text, style="List Number")

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph()

    # Title
    title = doc.add_heading("How the Unit-wise Scheduler Works", 0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = sub.add_run("A plain-language guide for leadership and manufacturing stakeholders")
    r.italic = True
    r.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.add_run(f"CMF Digitalization · PPS Machine Scheduling\n").bold = True
    meta.add_run(f"Document date: {date.today().strftime('%d %B %Y')}\n")
    meta.add_run(
        "Purpose: Explain how unit-wise planning works behind the scenes — "
        "without system or coding detail.\n"
    )

    # 1
    doc.add_heading("1. Why we built this", level=1)
    para(
        "Our shop is high-mix, low-volume: many different parts, relatively small "
        "quantities, and work that moves through several machines."
    )
    para(
        "Traditional planning often treats a whole batch as one block — for example, "
        "“finish all 10 pieces on Operation 10 before anyone starts Operation 20.” "
        "On the real floor, that is not always how work flows. When the first few pieces "
        "are approved, the next operation can often begin on those pieces while the "
        "rest of the batch is still finishing upstream."
    )
    para(
        "The Unit-wise Scheduler was built to plan at the level of each individual piece "
        "(each unit), so the plan matches that reality more closely — especially when "
        "partial quantities are released to the next operation."
    )

    # 2
    doc.add_heading("2. Two ways of looking at the same shop", level=1)
    para(
        "We did not replace the existing batch schedule. We added a second view that "
        "runs alongside it."
    )
    add_table(
        ["View", "What it shows", "Best for"],
        [
            [
                "Batch plan (Planned / Dynamic)",
                "Whole remaining quantity as larger blocks on machines",
                "Capacity overview and familiar batch planning",
            ],
            [
                "Unit-wise plan",
                "Each piece as its own bar (Unit 1, Unit 2, Unit 3…)",
                "Seeing flow, overlap, and remaining pieces after partial progress",
            ],
        ],
    )
    para(
        "Think of it like two lenses on the same factory: one shows batches; the other "
        "shows pieces moving through the route."
    )

    # 3
    doc.add_heading("3. The simple idea behind the hood", level=1)
    para("For every active in-house part, the unit-wise engine does this:")
    numbered("Take the quantity of the part (for example, 10 pieces).")
    numbered("Walk through each manufacturing operation in order (Op 10 → Op 20 → …).")
    numbered(
        "For pieces already finished and approved on an operation, do not plan them again "
        "on that operation."
    )
    numbered(
        "For pieces still remaining, place each one on a machine at the earliest sensible "
        "time, respecting shop rules (see next section)."
    )
    numbered(
        "When a piece finishes one operation, it becomes free to start the next — "
        "even if later pieces are still behind. That is the “pipeline” effect."
    )

    para("Example (quantity 10, two operations):", bold=True)
    bullet("Unit 1 finishes Op 10 and can move to Op 20.")
    bullet("Meanwhile Units 2–10 may still be on Op 10.")
    bullet(
        "The plan shows that overlap clearly — instead of waiting for all 10 to finish "
        "Op 10 before any Op 20 appears."
    )

    # 4
    doc.add_heading("4. Rules the scheduler respects (shop reality)", level=1)
    para(
        "Behind the scenes, placement is not free-form. It follows the same kinds of "
        "constraints planners and supervisors already care about."
    )
    add_table(
        ["Rule", "What it means in practice"],
        [
            [
                "Only active work",
                "Only parts that are activated for manufacturing are included.",
            ],
            [
                "In-house only",
                "Outsourced work is not planned in this unit-wise view.",
            ],
            [
                "Part priority",
                "Higher-priority jobs are considered ahead of lower-priority ones.",
            ],
            [
                "Shift timings",
                "Work is placed inside configured working shifts, not through the night "
                "unless the calendar allows it.",
            ],
            [
                "Preferred machine",
                "If a job is already tied to a machine (active job or existing plan), "
                "unit-wise stays on that machine instead of jumping to a sister machine.",
            ],
            [
                "Busy machines",
                "A machine with an operator currently running a job is not overwritten "
                "in the past; new plan starts from “now.”",
            ],
            [
                "Breakdown / machine off",
                "If a machine is marked unavailable, work waits until it is available again.",
            ],
            [
                "Setup vs cycle",
                "Full setup is charged when a new run starts; continuation and rework "
                "pieces typically use cycle time only (no repeated setup).",
            ],
            [
                "Progress already made",
                "Approved pieces are treated as done on that operation and are not "
                "re-planned there.",
            ],
        ],
    )

    # 5
    doc.add_heading("5. How this relates to the shop floor (approvals)", level=1)
    para(
        "The plan and the shop floor work together, but they have different jobs:"
    )
    bullet(
        "Unit-wise plan = “What is the best picture of remaining piece-by-piece work?”"
    )
    bullet(
        "Job cards / approvals = “What is the operator actually allowed to produce next?”"
    )
    para(
        "When supervisors approve partial quantity on Operation N, the next operation "
        "can unlock for that released quantity on the floor. When the unit-wise plan "
        "is refreshed, finished pieces drop out of that operation and the remaining "
        "pieces are re-laid on the timeline."
    )
    para(
        "So the scheduler is continuously aiming to stay aligned with real progress — "
        "not a one-time static picture."
    )

    # 6
    doc.add_heading("6. Two planning modes: Standard and Advanced search", level=1)
    para("There are two ways to build the unit-wise plan:")

    doc.add_heading("6.1 Standard (day-to-day)", level=2)
    para(
        "The standard method places remaining pieces in a practical order: "
        "follow part priority, follow the route, put each piece on the preferred or "
        "earliest suitable machine, and move forward. This is fast, explainable, and "
        "suitable for regular use."
    )

    doc.add_heading("6.2 Advanced search (optional)", level=2)
    para(
        "Sometimes the order of pieces across shared machines can be improved — "
        "for example to reduce waiting, reduce extra setups, finish urgent work sooner, "
        "or improve throughput."
    )
    para(
        "The advanced mode tries many alternative piece sequences under the same shop "
        "rules, then keeps a new plan only if it is better than the standard plan on "
        "the agreed goals. If it finds no improvement, the system keeps the standard plan."
    )
    para(
        "In short: Advanced search is an optional “try to do better” step — not a "
        "replacement for shop rules."
    )

    doc.add_heading("6.3 What “better” means", level=2)
    para("When comparing plans, we look for outcomes management typically cares about:")
    bullet("Higher effective use of machines")
    bullet("Fewer unnecessary setups")
    bullet("Less waiting and idle gaps")
    bullet("Shorter time for pieces to move through the route (faster flow)")
    bullet("Higher throughput (more pieces completed per hour of schedule window)")
    bullet("Less lateness against due dates (when due dates exist)")
    bullet("Respect for part priority")

    # 7
    doc.add_heading("7. How we prove it helps — side-by-side comparison", level=1)
    para(
        "For any chosen part, the system can show the same performance measures for:"
    )
    numbered("Batch Planned schedule")
    numbered("Batch Dynamic schedule")
    numbered("Unit-wise schedule")
    para("Typical measures include:")
    add_table(
        ["Measure", "Plain meaning"],
        [
            ["Completion window", "How long from first start to last finish"],
            ["Flow time", "How long a piece stays “in the system”"],
            ["Waiting between operations", "Idle gaps between Op 10 and Op 20, etc."],
            ["Machine utilization", "How densely the machine is used in that window"],
            ["Machine idle", "Empty gaps on the machine inside that window"],
            ["Throughput", "Pieces finished per hour of that window"],
            ["Lateness / early finish", "Against the order due date, if set"],
        ],
    )
    para(
        "This comparison is the management proof point: we can see whether unit-wise "
        "improves flow and delivery behaviour versus the batch views — on real parts, "
        "not only in theory."
    )

    # 8
    doc.add_heading("8. What you see in the system (for planners)", level=1)
    bullet("On the Planned Schedule screen, switch between Batch-wise and Unit-wise.")
    bullet("In Unit-wise, each bar is a piece (for example U-001, U-002).")
    bullet("Choose Standard rebuild or Advanced search, then refresh the plan.")
    bullet(
        "Open the comparison panel to read Planned vs Dynamic vs Unit-wise metrics "
        "for a selected part."
    )

    # 9
    doc.add_heading("9. What success looks like", level=1)
    para("Unit-wise is working as intended when:")
    bullet(
        "Planners can see remaining pieces clearly after partial production progress"
    )
    bullet(
        "The plan respects shifts, preferred machines, priorities, and machine unavailability"
    )
    bullet(
        "On multi-operation work, pieces can overlap across operations instead of "
        "waiting for the full batch"
    )
    bullet(
        "Comparison against Planned and Dynamic shows a clear story on flow, waiting, "
        "utilization, or on-time performance — especially on suitable pilot parts"
    )

    para("It is less useful when:", bold=True)
    bullet("Only one operation is left with a tiny remaining quantity")
    bullet("Due dates are so far away that lateness measures say little")
    bullet(
        "The question is only “fill every machine to 100% plant-wide” — this view is "
        "part-focused planning, not full plant OEE accounting"
    )

    # 10
    doc.add_heading("10. One-page story for leadership", level=1)
    para(
        "We still plan in batches for the familiar capacity picture. In parallel, "
        "we now also plan piece-by-piece so the schedule matches how a high-mix shop "
        "actually releases work: finish and approve some pieces, start the next "
        "operation on those pieces, and keep the rest moving upstream."
    )
    para(
        "The engine follows real shop rules — active jobs only, priorities, shifts, "
        "preferred machines, breakdowns, and progress already made. An optional "
        "advanced search can try to improve that plan against utilization, setups, "
        "waiting, flow, throughput, and lateness — but only keeps the result if it "
        "is truly better. Side-by-side measures against Planned and Dynamic schedules "
        "let us prove the value on real parts."
    )

    doc.add_paragraph()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = end.add_run("— End of document —")
    er.italic = True

    footer = doc.add_paragraph()
    footer.add_run(
        "Prepared for internal sharing with managers and manufacturing stakeholders. "
        "No system configuration or developer instructions are required to read this document."
    ).italic = True

    out = Path(__file__).resolve().parent / "UNIT_WISE_SCHEDULER_HOW_IT_WORKS.docx"
    doc.save(out)
    print(str(out))


if __name__ == "__main__":
    main()
