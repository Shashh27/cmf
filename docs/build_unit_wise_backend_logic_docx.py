"""Generate UNIT_WISE_SCHEDULER_BACKEND_BUSINESS_LOGIC.docx"""

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt


def main() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def para(text: str, bold: bool = False):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        return p

    def bullet(text: str):
        doc.add_paragraph(text, style="List Bullet")

    def numbered(text: str):
        doc.add_paragraph(text, style="List Number")

    def add_table(headers, rows):
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                table.rows[ri + 1].cells[ci].text = str(val)
        doc.add_paragraph()

    # ── Title ──
    doc.add_heading("Unit-wise Scheduler — Backend & Business Logic", 0)
    sub = doc.add_paragraph()
    r = sub.add_run("Complete technical reference for PPS / CMF Digitalization")
    r.bold = True
    r.font.size = Pt(13)

    meta = doc.add_paragraph()
    meta.add_run(f"Date: {date.today().isoformat()}\n").italic = True
    meta.add_run("Audience: Backend engineers, manufacturing systems, PPS stakeholders\n").italic = True
    meta.add_run(
        "Code roots: backend/cmf/unit_wise_scheduler.py, unit_wise_ga_research.py, "
        "unit_wise_compare.py, routers/unit_wise.py\n"
    ).italic = True

    para(
        "This document describes the full backend design and business rules of the "
        "Unit-wise Scheduler: why it exists, how it decides times and machines, how it "
        "interacts with production unlocks, how GA optimizes it, and how KPIs are compared "
        "against batch Planned and Dynamic schedules."
    )

    # 1 Business purpose
    doc.add_heading("1. Business Purpose", level=1)
    para(
        "CMF is a high-mix low-volume (HMLV) shop. Batch scheduling plans whole remaining "
        "quantities as contiguous blocks. That is useful for capacity views, but it does not "
        "reflect unit-level flow when partial quantities are approved and the next operation "
        "can start on released pieces."
    )
    para("Unit-wise scheduling exists to:")
    bullet("Plan each physical unit (1..qty) through each schedulable operation")
    bullet("Allow pipeline overlap (Unit 1 on Op20 while Unit 2 still on Op10)")
    bullet("Re-plan remaining units after production approvals / rework")
    bullet("Keep batch Planned and Dynamic Gantt unchanged (dual view)")
    bullet("Optionally search better unit orders with a research-grade GA")

    doc.add_heading("1.1 What unit-wise is NOT", level=2)
    bullet("It does not replace batch Planned (planned_schedule_items)")
    bullet("It does not replace batch Dynamic (rescheduling_items / algorithm.py generate)")
    bullet("It does not by itself unlock job cards (shop unlock remains production_log_helpers)")
    bullet("It does not invent due dates; tardiness uses oms.orders.due_date when present")

    # 2 Dual view
    doc.add_heading("2. Dual Schedule Model", level=1)
    add_table(
        ["View", "Table", "Grain", "Owner"],
        [
            [
                "Batch Planned",
                "scheduling.planned_schedule_items",
                "Operation block (batch qty)",
                "Generate / history active version",
            ],
            [
                "Batch Dynamic",
                "scheduling.rescheduling_items",
                "Operation block (remaining qty)",
                "algorithm.SchedulerEngine / generate-schedule",
            ],
            [
                "Unit-wise",
                "scheduling.unit_schedule_items",
                "One unit × one op segment",
                "unit_wise_scheduler / GA research",
            ],
        ],
    )
    para(
        "Business rule: Unit-wise rebuild never INSERT/UPDATE/DELETE batch rows. It only "
        "rewrites unit_schedule_items for the selected active scope."
    )

    # 3 Data model
    doc.add_heading("3. Data Model — UnitScheduleItem", level=1)
    para("Schema: scheduling.unit_schedule_items (model: DB.models.scheduling.UnitScheduleItem)")
    add_table(
        ["Column", "Meaning"],
        [
            ["order_id / order_number", "Sale order identity"],
            ["part_id / part_number", "Part identity"],
            ["unit_index", "Physical unit number 1..part.qty"],
            ["operation_id / operation_number", "Routing step"],
            ["machine_id", "Assigned machine for this segment"],
            ["start_time / end_time", "Planned segment window (may split across shifts)"],
            ["status", "Typically unit_scheduled"],
            ["schedule_version", "Increments on each rebuild"],
            ["source", "greedy or ga_research (what produced the row)"],
            ["created_at", "Insert timestamp"],
        ],
    )
    para(
        "One logical unit-op placement may produce multiple rows if duration spans multiple "
        "shift windows (segmented placement)."
    )

    # 4 Scope
    doc.add_heading("4. Scope Selection (Which Work Is Planned)", level=1)
    para("Function: _load_active_scope(db, part_id?, order_id?)")
    bullet("Join PartScheduleStatus + Part + Order + OrderPartPriority")
    bullet("Filter: PartScheduleStatus.status == 'active'")
    bullet("Filter: Part.type_id == IN_HOUSE (1) — outsourced not planned here")
    bullet("Require part.qty > 0")
    bullet("Optional filters: single part_id and/or order_id")
    bullet("Sort: (order_part_priority ASC, part.id ASC) — lower priority number = higher urgency")
    para(
        "Business meaning: only activated in-house parts enter the unit-wise plan. "
        "Inactive / not-yet-activated parts are ignored until activation."
    )

    # 5 Schedulable ops
    doc.add_heading("5. Schedulable Operations", level=1)
    para("Function: _schedulable_operations → filters via is_schedulable_operation")
    bullet("Operations must be in-house / schedulable per production_log_helpers rules")
    bullet("Sorted by numeric operation_number when possible")
    para(
        "Non-schedulable ops (e.g. certain outsourced or excluded types) never create "
        "unit-wise segments."
    )

    # 6 Duration
    doc.add_heading("6. Duration & Setup Business Rules", level=1)
    para("Function: _duration(operation, skip_setup)")
    para("duration_seconds = (0 if skip_setup else setup_time) + cycle_time")
    para("If total ≤ 0 → force 60 seconds (avoid zero-width Gantt bars)")
    doc.add_heading("6.1 When setup is skipped", level=2)
    bullet("Same machine just ran same (part_id, operation_id) — continuation of a run")
    bullet("Operation already has approved qty > 0 — setup already done on the floor")
    bullet("Rework slot: first rework_due remaining units after review — cycle only")
    para(
        "Business meaning: avoid charging setup again for rework pieces or mid-run "
        "unit continuations."
    )

    # 7 Machine selection
    doc.add_heading("7. Machine Selection Business Rules", level=1)
    doc.add_heading("7.1 Candidate machines", level=2)
    bullet("Primary: all machines in operation.workcenter_id")
    bullet(
        "If preferred machine is outside WC list, it is prepended so pin can still apply"
    )

    doc.add_heading("7.2 Preferred machine resolution", level=2)
    para("Function: _preferred_machine_id(db, operation, order_id)")
    numbered("1. Active ProductionLog (operator_status=inprogress) on this operation → its machine")
    numbered("2. Else latest Rescheduling row for this operation (optionally filtered by order)")
    numbered("3. Else operation.machine_id from routing master")
    para(
        "Business meaning: unit-wise follows the machine the shop already committed to, "
        "reducing sibling WC jumps (e.g. Magerle vs HMT-500)."
    )

    doc.add_heading("7.3 Pick logic", level=2)
    para("Function: _pick_machine(machines, machine_free, ready, preferred_id)")
    bullet(
        "If UNIT_WISE_PIN_PREFERRED=true (default) and preferred_id set → always that machine"
    )
    bullet(
        "Else choose machine with earliest start = max(machine_free[m], unit_ready)"
    )

    doc.add_heading("7.4 Freeze in-progress machines", level=2)
    para("Function: _freeze_active_machines")
    bullet(
        "Any machine with an inprogress ProductionLog gets machine_free raised to now"
    )
    para(
        "Business meaning: do not schedule other unit work into the past or over an "
        "active job card."
    )

    # 8 Time placement
    doc.add_heading("8. Time Placement — Shifts & Breakdowns", level=1)
    doc.add_heading("8.1 Start candidate", level=2)
    para("start_candidate = max(unit_ready[u], machine_free[machine], now)")
    para(
        "Part activation / schedule start is also snapped into shift via "
        "SchedulerEngine.adjust_to_shift (fallback: default 08:30–17:00)."
    )

    doc.add_heading("8.2 Shift splitting", level=2)
    para("Function: _place_within_shifts_engine(engine, machine_id, start, duration)")
    bullet("Uses real ShiftHoursConfiguration via SchedulerEngine")
    bullet("Honours machine OFF windows via engine._machine_next_available")
    bullet("Splits long jobs across remaining shift capacity → multiple segments")
    bullet("Falls back to default daily windows if engine helpers fail")

    doc.add_heading("8.3 Breakdown / OFF (business behaviour)", level=2)
    bullet(
        "If machine is OFF at candidate time, placement waits until available_to "
        "(or next shift)"
    )
    bullet("If permanently OFF (available_to NULL), machine cannot be used")
    bullet(
        "GA decode additionally calls _machine_next_available explicitly and may try "
        "an alternate WC machine when pin is off"
    )
    para(
        "Business meaning: unit-wise does not schedule through known breakdowns; "
        "work is deferred to the next available ON window."
    )

    # 9 Greedy pipeline
    doc.add_heading("9. Greedy Pipeline Algorithm (Core Business Logic)", level=1)
    para("Entry: simulate_unit_plan(...) / rebuild_unit_schedule(optimizer='greedy')")
    para("Conceptual loop:")
    numbered("Load active scope (priority ordered)")
    numbered("Initialize machine_free, freeze in-progress machines, create SchedulerEngine")
    numbered("For each part in scope:")
    bullet("    Set unit_ready[1..qty] = part_start (activation/now, shift-adjusted)")
    bullet("    For each schedulable operation in route order:")
    bullet("        approved = total_approved_for_operation")
    bullet("        For units 1..approved: advance unit_ready using actual_end (no new rows)")
    bullet("        remaining_units = approved+1 .. qty (optionally reordered by GA/policy)")
    bullet("        For each remaining unit:")
    bullet("            pick machine (pin / earliest)")
    bullet("            decide skip_setup")
    bullet("            place duration into shifts → append segment dicts/rows")
    bullet("            machine_free[m] = last_end; unit_ready[u] = last_end")
    para(
        "Pipeline effect: after Unit 1 finishes Op10, unit_ready[1] allows Op20 while "
        "Unit 2 may still be waiting/running Op10 on the same or another timeline."
    )

    doc.add_heading("9.1 Production coupling", level=2)
    add_table(
        ["Shop event", "Unit-wise reaction"],
        [
            [
                "Partial approve on Op N",
                "Those unit indexes skipped on Op N; downstream remaining units re-planned on rebuild",
            ],
            [
                "Actual end times on logs",
                "Approved units’ ready times pushed to actual completion",
            ],
            [
                "Rework qty on latest reviewed log",
                "Next rework_due remaining slots are cycle-only (no re-setup)",
            ],
            [
                "Job card inprogress on machine",
                "That machine frozen to now for new unit placements",
            ],
        ],
    )
    para(
        "Important: unit-wise plan visibility is not the same as job-card unlock. "
        "Unlock quantity caps remain enforced by production_log_helpers / machine_scheduling "
        "APIs (shop hard unlock)."
    )

    # 10 Rebuild
    doc.add_heading("10. Rebuild Persistence Flow", level=1)
    para("Function: rebuild_unit_schedule(db, part_id?, order_id?, commit, optimizer)")
    numbered("Abort if UNIT_WISE_SCHEDULE_ENABLED is false")
    numbered("Load scope; if empty return success with 0 rows")
    numbered("DELETE existing unit_schedule_items for part_ids in scope")
    numbered("schedule_version = max(existing versions)+1 (global max then +1)")
    numbered("Build plan via greedy simulate OR research GA")
    numbered("INSERT UnitScheduleItem rows with that version and source")
    numbered("commit/flush; return counts, version, makespan, ga meta")
    para(
        "Business meaning: each rebuild is a full refresh for the scoped parts — "
        "no merge of stale unit bars."
    )

    # 11 GA
    doc.add_heading("11. Research GA Business Logic (Optional Optimizer)", level=1)
    para("Module: unit_wise_ga_research.py — used when optimizer is ga or ga_research")

    doc.add_heading("11.1 Why GA exists", level=2)
    para(
        "Greedy places remaining units in index order with earliest-start machine policy. "
        "That can be suboptimal for multi-part contention, setup sequences, and due-driven "
        "trade-offs. GA searches alternative activity orders under the same hard constraints."
    )

    doc.add_heading("11.2 Activity model", level=2)
    bullet("One activity = one unfinished (part, unit_index, operation)")
    bullet("Predecessor = previous unfinished op of the same unit (routing precedence)")
    bullet("Carries part_priority, preferred_id, allowed WC machines, rework_slot flag")

    doc.add_heading("11.3 Chromosome & decode", level=2)
    bullet("π = permutation of all activities (operation-sequence / priority list)")
    bullet("Optional machine genes when pin is off / no preferred")
    bullet(
        "Decode: repeatedly schedule the first precedence-feasible activity in π "
        "(semi-active priority list)"
    )
    bullet("Placement uses same shift/breakdown/setup rules as greedy")

    doc.add_heading("11.4 Fitness (business end goals)", level=2)
    para("Minimize cost; fitness = −cost. Terms:")
    add_table(
        ["Business goal", "Term"],
        [
            ["Finish sooner", "w_makespan · Cmax"],
            ["Faster unit flow", "w_flow · mean flow time"],
            ["Less waiting between ops", "w_wait · mean waiting"],
            ["On-time delivery", "w_tard · mean tardiness"],
            ["Fewer changeovers", "w_setup · setup_count"],
            ["Less lost capacity", "w_idle · idle hours"],
            ["Higher utilization", "w_util_gap · (100−util)/100"],
            ["More output rate", "− w_throughput · units/hour"],
            ["Respect part priority", "w_priority · priority inversions"],
        ],
    )
    para(
        "Acceptance gate: persist GA plan only if scalar fitness strictly beats greedy; "
        "otherwise keep greedy (source=greedy)."
    )

    doc.add_heading("11.5 Operators & experiment design", level=2)
    bullet("Tournament selection, Order Crossover (OX), swap + inversion mutation, elitism")
    bullet("Population seeded with natural order + priority-sorted order")
    bullet("Multi-run with seeds → stats: cmax best/mean/std/worst + convergence series")
    bullet("Large instances (>120 activities) auto-scale pop/gens/runs down")

    # 12 Compare
    doc.add_heading("12. Compare Engine Business Logic", level=1)
    para("Module: unit_wise_compare.py — GET /scheduling/unit-wise/compare")
    para("For a part (or each part of an order), compute side-by-side:")
    bullet("batch_planned — active planned_schedule_items")
    bullet("batch_dynamic — rescheduling_items (scheduled/rescheduled)")
    bullet("unit_wise — latest unit_schedule_items")

    doc.add_heading("12.1 KPI definitions", level=2)
    add_table(
        ["KPI", "Formula / meaning"],
        [
            ["Makespan", "latest_end − earliest_start of segments"],
            [
                "Unit flow",
                "For each unit: last segment end − first segment start",
            ],
            [
                "Batch flow proxy",
                "Treat batch block as one job ≈ makespan",
            ],
            [
                "Waiting",
                "Sum of gaps between consecutive op blocks (next start − prev end)",
            ],
            [
                "Machine busy",
                "Merged busy time on that machine for this part",
            ],
            [
                "Machine idle",
                "machine span − busy (gaps inside the part’s window on that machine)",
            ],
            [
                "Utilization %",
                "busy / span × 100 (this part only, not plant-wide)",
            ],
            [
                "Throughput",
                "unfinished_units / makespan (aligned to unit-wise count when present)",
            ],
            [
                "Tardiness",
                "max(0, finish − due_date)",
            ],
            [
                "Earliness",
                "max(0, due_date − finish)",
            ],
        ],
    )
    para(
        "UI surfaces Planned | Dynamic | Unit-wise with Δ vs Dynamic and Δ vs Planned "
        "so stakeholders can judge whether unit-wise outperforms batch views."
    )

    # 13 API
    doc.add_heading("13. REST API Surface", level=1)
    add_table(
        ["Method", "Path", "Business action"],
        [
            [
                "POST",
                "/api/v1/scheduling/unit-wise/rebuild",
                "Rebuild greedy or GA plan for scope (part_id/order_id optional)",
            ],
            [
                "GET",
                "/api/v1/scheduling/unit-wise",
                "List latest (or all) unit segments; filters part/order/machine",
            ],
            [
                "GET",
                "/api/v1/scheduling/unit-wise/compare",
                "KPI compare Planned / Dynamic / Unit-wise",
            ],
            [
                "GET",
                "/api/v1/scheduling/unit-wise/parts/{part_id}",
                "Segments for one part",
            ],
            [
                "GET",
                "/api/v1/scheduling/unit-wise/machines/{machine_id}",
                "Segments for one machine",
            ],
        ],
    )
    para("Rebuild body example:")
    para('{ "optimizer": "ga", "part_id": 1661 }')
    para("optimizer ∈ { greedy, ga, ga_research } — ga and ga_research use research engine.")

    # 14 Config
    doc.add_heading("14. Configuration (Environment)", level=1)
    add_table(
        ["Variable", "Default", "Business effect"],
        [
            ["UNIT_WISE_SCHEDULE_ENABLED", "true", "Master on/off for APIs & rebuild"],
            ["UNIT_WISE_OPTIMIZER", "greedy", "Default optimizer if body omits it"],
            ["UNIT_WISE_PIN_PREFERRED", "true", "Hard pin preferred machine"],
            ["UNIT_WISE_GA_ENABLED", "true", "Allow GA path"],
            ["UNIT_WISE_GA_POPULATION", "40", "GA population size"],
            ["UNIT_WISE_GA_GENERATIONS", "60", "Generations per run"],
            ["UNIT_WISE_GA_RUNS", "3", "Independent experimental runs"],
            ["UNIT_WISE_GA_SEED", "42", "Base RNG seed"],
            ["UNIT_WISE_GA_W_*", "(see code)", "Fitness weights for end goals"],
        ],
    )

    # 15 Module map
    doc.add_heading("15. Backend Module Map", level=1)
    add_table(
        ["Module", "Responsibility"],
        [
            ["unit_wise_scheduler.py", "Scope, greedy simulate, shift place, rebuild, list"],
            ["unit_wise_ga_research.py", "Activities, OX GA, multi-obj fitness, multi-run"],
            ["unit_wise_ga.py", "Delegate entry to research GA"],
            ["unit_wise_compare.py", "Planned/Dynamic/Unit KPI engine"],
            ["routers/unit_wise.py", "HTTP API"],
            ["DB/models/scheduling.py", "UnitScheduleItem ORM"],
            ["migrations/create_unit_schedule_items.py", "Schema create/add columns"],
            ["production_log_helpers.py", "Approved qty, schedulable ops, unlock helpers"],
            ["algorithm.SchedulerEngine", "Shift calendar + machine OFF resolution"],
        ],
    )

    # 16 End-to-end
    doc.add_heading("16. End-to-End Business Flow", level=1)
    numbered("PEC / coordinator activates in-house part (PartScheduleStatus=active)")
    numbered("Batch Planned/Dynamic may already exist from generate-schedule")
    numbered("User (or automatic hook after schedule generate / production review) calls unit-wise rebuild")
    numbered("System deletes old unit rows for scope and inserts new version")
    numbered("Planners view Unit-wise Gantt (U-001 bars) separately from batch")
    numbered("Operators run job cards; partial approve unlocks downstream qty (shop rules)")
    numbered("Rebuild again → approved units disappear from that op; remainder re-pipelined")
    numbered("Phase 3 compare proves Planned vs Dynamic vs Unit-wise KPIs")

    doc.add_heading("16.1 Automatic rebuild hooks", level=2)
    bullet(
        "machine_scheduling generate-schedule path can call rebuild_unit_schedule after batch generate"
    )
    bullet(
        "production_logs review/approve path can call rebuild_unit_schedule for affected part "
        "so remaining units re-pipeline after approvals/rework"
    )
    para(
        "Hooks are best-effort and gated by unit_wise_enabled(); failures are logged and must "
        "not break the primary batch/production transaction."
    )

    # 17 Limitations
    doc.add_heading("17. Business Limitations & Caveats", level=1)
    bullet(
        "Strongest value on multi-op unfinished work with partial qty release; weak on "
        "single remaining op remnants"
    )
    bullet("Far due dates make tardiness/earliness KPIs uninformative")
    bullet(
        "With preferred pin ON, GA cannot freely load sibling machines — by design for shop stability"
    )
    bullet(
        "Unit-wise bars may show look-ahead beyond shop unlock unless/until plan≡unlock is adopted"
    )
    bullet(
        "Utilization in compare is part-local (busy/span for that part), not plant OEE"
    )

    # 18 Glossary
    doc.add_heading("18. Glossary", level=1)
    add_table(
        ["Term", "Meaning in this system"],
        [
            ["Unit", "One piece of part.qty (unit_index)"],
            ["Segment", "One continuous time interval on a machine for a unit-op"],
            ["Pipeline", "Downstream op starts for unit u before later units finish upstream"],
            ["Pin", "Force preferred machine instead of earliest sibling"],
            ["Freeze", "Do not schedule before now on machines with active jobs"],
            ["Greedy", "Earliest-start list scheduling without search"],
            ["GA Research", "OX permutation search with multi-objective cost"],
            ["Makespan", "Schedule window length for the compared segments"],
            ["Flow time", "Time a unit spends from first start to last end"],
        ],
    )

    doc.add_heading("19. Conclusion", level=1)
    para(
        "The Unit-wise Scheduler is a full backend subsystem: scoped to active in-house "
        "parts, driven by routing and production progress, placed with shift/breakdown/"
        "pin/freeze rules, persisted independently of batch schedules, optionally optimized "
        "by research GA, and measured against Planned and Dynamic via explicit KPIs. "
        "It supports HMLV partial-flow planning while leaving shop unlock and batch engines intact."
    )
    para("— End of backend & business logic reference —", bold=True)

    out = (
        Path(__file__).resolve().parent
        / "UNIT_WISE_SCHEDULER_BACKEND_BUSINESS_LOGIC.docx"
    )
    doc.save(out)
    print(str(out))


if __name__ == "__main__":
    main()
