"""
Purchase Request Service — generates populated .docx from Purchase Requisition templates.
"""
import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from sqlalchemy.orm import Session, joinedload

from DB.models.inventory import (
    RawMaterialStock as RawMaterialStockModel,
    Vendors as VendorsModel,
)
from DB.models.oms import Order as OrderModel
from DB.models.access_control import AccessUser as AccessUserModel

# Template paths for purchase request
ASSETS_DIR = os.path.join(os.path.dirname(__file__), "..", "assets")
TEMPLATE_UP_25000 = os.path.join(ASSETS_DIR, "FormPurchaseReq-Up-25000.docx")
TEMPLATE_25000_TO_50000 = os.path.join(ASSETS_DIR, "FormPurchaseReq-25000to50000.docx")
TEMPLATE_MORE_THAN_50000 = os.path.join(ASSETS_DIR, "FormPurchaseReqMorethan50000.docx")


def _set_cell_text(cell, text: str):
    """Set cell text preserving existing paragraph/run formatting where possible."""
    try:
        if not cell or not text:
            return
        if not cell.paragraphs:
            cell.text = text
            return
        p = cell.paragraphs[0]
        if p.runs:
            p.runs[0].text = text
            for run in p.runs[1:]:
                run.text = ""
            for extra_p in cell.paragraphs[1:]:
                for run in extra_p.runs:
                    run.text = ""
        else:
            p.text = text
    except Exception as e:
        print(f"Error in _set_cell_text: {e}")


def _set_cell_text_with_font(cell, text: str, font_size_pt: int | None = None):
    """Set cell text and optionally apply a specific font size to the inserted text."""
    try:
        _set_cell_text(cell, text)
        if font_size_pt is not None:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(font_size_pt)
    except Exception as e:
        print(f"Error in _set_cell_text_with_font: {e}")


def _set_paragraph_text(para, text: str):
    """Replace paragraph text while preserving alignment and formatting."""
    try:
        if not para or not text:
            return
        original_alignment = para.alignment
        for run in list(para.runs):
            run._element.getparent().remove(run._element)
        if para.runs:
            new_run = para.add_run(text)
            if para.runs[0]:
                new_run.font.name = para.runs[0].font.name
                new_run.font.size = para.runs[0].font.size
                new_run.bold = para.runs[0].bold
                new_run.italic = para.runs[0].italic
        else:
            para.add_run(text)
        para.alignment = original_alignment
    except Exception as e:
        print(f"Error in _set_paragraph_text: {e}")


def _get_user_name(user_id: int, db: Session) -> str:
    """Resolve user_id to actual name."""
    if not user_id:
        return ""
    user = db.query(AccessUserModel).filter(AccessUserModel.id == user_id).first()
    if user:
        return user.user_name or user.name or ""
    return ""


def _get_order_details(order_id: int, db: Session) -> dict:
    """Fetch order details from database by ID."""
    if not order_id:
        return None
    try:
        order = db.query(OrderModel).filter(OrderModel.id == order_id).first()
        if order:
            return {
                "id": order.id,
                "order_number": order.order_number,
                "project_name": getattr(order, 'project_name', '') or ""
            }
    except Exception as e:
        print(f"Error fetching order details: {e}")
    return None


def generate_purchase_request_docx(
    stock_id: int,
    template_type: str,
    data: dict,
    db: Session
):
    """
    Generate a populated Purchase Requisition .docx for order material.
    Accepts edited data from frontend and fills in the template.
    
    Args:
        stock_id: ID of the raw material stock
        template_type: Template type (auto, up_25000, 25000_to_50000, more_than_50000)
        data: Dictionary containing edited form data
        db: Database session
    
    Returns:
        Tuple of (buffer, filename) for the generated document
    """
    # Get stock with all related data
    stock = db.query(RawMaterialStockModel).options(
        joinedload(RawMaterialStockModel.material),
        joinedload(RawMaterialStockModel.source_order),
        joinedload(RawMaterialStockModel.vendor),
    ).filter(RawMaterialStockModel.id == stock_id).first()
    
    if not stock:
        raise ValueError("Stock not found")
    
    # Check if this is a grouped order and fetch all records in the group
    grouped_stocks = []
    try:
        if stock.merge_group_id:
            grouped_stocks = db.query(RawMaterialStockModel).options(
                joinedload(RawMaterialStockModel.material),
                joinedload(RawMaterialStockModel.source_order),
            ).filter(
                RawMaterialStockModel.merge_group_id == stock.merge_group_id,
                RawMaterialStockModel.source_type == "order"
            ).all()
        else:
            grouped_stocks = [stock]
    except Exception as e:
        print(f"Error fetching grouped stocks: {e}")
        grouped_stocks = [stock]
    
    # Determine template based on cost or explicit selection
    cost = data.get("cost", stock.final_cost or stock.estimated_cost or 0)
    
    if template_type == "auto":
        if cost <= 25000:
            template_path = TEMPLATE_UP_25000
            case_num = 1
        elif cost <= 50000:
            template_path = TEMPLATE_25000_TO_50000
            case_num = 2
        else:
            template_path = TEMPLATE_MORE_THAN_50000
            case_num = 3
    elif template_type == "up_25000":
        template_path = TEMPLATE_UP_25000
        case_num = 1
    elif template_type == "25000_to_50000":
        template_path = TEMPLATE_25000_TO_50000
        case_num = 2
    elif template_type == "more_than_50000":
        template_path = TEMPLATE_MORE_THAN_50000
        case_num = 3
    else:
        raise ValueError("Invalid template type")
    
    # Load template
    template_path = os.path.normpath(template_path)
    if not os.path.exists(template_path):
        raise ValueError(f"Docx template for Case {case_num} not found on server")
    
    try:
        doc = Document(template_path)
    except Exception as e:
        raise ValueError(f"Error loading template: {e}")
    
    # Prepare created date string
    created_date = stock.created_at.strftime("%d/%m/%Y") if stock.created_at else ""
    
    # Replace the existing दिनांक/Date: placeholder
    for para in doc.paragraphs:
        if "दिनांक/Date:" in para.text:
            new_text = para.text.replace("दिनांक/Date:", f"दिनांक/Date: {created_date}")
            _set_paragraph_text(para, new_text)
            break
    
    # Get project details from edited data
    project_number = data.get("project_number", "")
    project_name = data.get("project_name", "")
    
    # Insert project number into declaration
    if project_number:
        for para in doc.paragraphs:
            if "It is certified that the required item/s is/are will be utilized for the project" in para.text:
                for idx, run in enumerate(para.runs):
                    if "/Department" in run.text:
                        run.text = run.text.replace("/Department", f"/Department {project_number}({project_name})")
                        run.text = run.text.replace("_", "")
                        if idx + 1 < len(para.runs):
                            para.runs[idx + 1].text = para.runs[idx + 1].text.lstrip("_")
                        break
                break
    
    # Build material description from edited data
    # If grouped orders, show "attached in annexure 1"
    is_grouped = len(grouped_stocks) > 1
    material_desc = data.get("material_required", "")
    
    # Calculate total quantity for grouped orders
    total_quantity = 0
    if is_grouped:
        for gs in grouped_stocks:
            if gs.quantity:
                total_quantity += gs.quantity
    else:
        total_quantity = data.get("quantity", stock.quantity or 0)
    
    if is_grouped:
        material_desc += "\n(Attached in Annexure 1)"
        # Also show process type and form type for grouped orders
        if data.get("process_type"):
            material_desc += f"\nProcess Type: {data.get('process_type')}"
        if data.get("form_type"):
            material_desc += f"\nForm Type: {data.get('form_type')}"
    else:
        # Single order - show dimensions as before
        if data.get("form_type"):
            material_desc += f"\nForm: {data.get('form_type')}"
        
        # Add dimensions based on form type
        form_type = data.get("form_type", "")
        if form_type == "Round":
            if data.get("diameter"):
                material_desc += f"\nDiameter: {data.get('diameter')}mm"
            if data.get("length"):
                material_desc += f"\nLength: {data.get('length')}mm"
        elif form_type == "Square":
            if data.get("breadth"):
                material_desc += f"\nBreadth: {data.get('breadth')}mm"
            if data.get("height"):
                material_desc += f"\nHeight: {data.get('height')}mm"
            if data.get("length"):
                material_desc += f"\nLength: {data.get('length')}mm"
        elif form_type == "Pipe":
            if data.get("inner_diameter"):
                material_desc += f"\nInner Diameter: {data.get('inner_diameter')}mm"
            if data.get("outer_diameter"):
                material_desc += f"\nOuter Diameter: {data.get('outer_diameter')}mm"
            if data.get("length"):
                material_desc += f"\nLength: {data.get('length')}mm"
        else:
            # Fallback for other form types - show all available dimensions
            if data.get("diameter"):
                material_desc += f"\nDiameter: {data.get('diameter')}mm"
            if data.get("length"):
                material_desc += f"\nLength: {data.get('length')}mm"
            if data.get("breadth"):
                material_desc += f"\nBreadth: {data.get('breadth')}mm"
            if data.get("height"):
                material_desc += f"\nHeight: {data.get('height')}mm"
    
    # Fill Table 0 (main form table) with edited data
    if doc.tables and len(doc.tables) > 0:
        table = doc.tables[0]
        
        if case_num in (1, 2):
            # Row 0: Indenting Officer name + Designation
            try:
                if len(table.rows) > 0 and len(table.rows[0].cells) > 2:
                    officer_text = data.get("indenting_officer", "")
                    designation = data.get("designation", "")
                    if designation:
                        officer_text = f"{officer_text}\n{designation}"
                    _set_cell_text(table.rows[0].cells[2], officer_text)
            except Exception as e:
                print(f"Error setting row 0: {e}")
            
            # Row 1: Centre/Group
            try:
                if len(table.rows) > 1 and len(table.rows[1].cells) > 2:
                    _set_cell_text(table.rows[1].cells[2], data.get("centre_group", ""))
            except Exception as e:
                print(f"Error setting row 1: {e}")
            
            # Row 2: Material/Service Required + Quantity
            try:
                material_text = material_desc
                if total_quantity:
                    material_text += f"\n{total_quantity}"
                if len(table.rows) > 2 and len(table.rows[2].cells) > 2:
                    _set_cell_text(table.rows[2].cells[2], material_text)
            except Exception as e:
                print(f"Error setting row 2: {e}")
            
            # Row 3: Project details
            try:
                project_text_parts = []
                if project_number:
                    project_text_parts.append(f"a) {project_number} - {project_name}")
                if data.get("budget_head"):
                    project_text_parts.append(f"b) {data.get('budget_head')}")
                project_text_parts.append(f"c) ₹{cost:,.2f}")
                if len(table.rows) > 3 and len(table.rows[3].cells) > 2:
                    _set_cell_text(table.rows[3].cells[2], "\n".join(project_text_parts))
            except Exception as e:
                print(f"Error setting row 3: {e}")
            
            # Row 4: Source of supply
            try:
                if len(table.rows) > 4 and len(table.rows[4].cells) > 2:
                    _set_cell_text(table.rows[4].cells[2], data.get("source_of_supply", ""))
            except Exception as e:
                print(f"Error setting row 4: {e}")
            
            # Row 5: Processed by (for Case 1 and Case 2)
            try:
                if case_num == 1:
                    processed_by = data.get("processed_by", "purchase")
                    if processed_by == "purchase":
                        processed_text = "☑ क्रय / Purchase      □ इंडेंटर / Indenter"
                    elif processed_by == "indenter":
                        processed_text = "□ क्रय / Purchase      ☑ इंडेंटर / Indenter"
                    else:
                        processed_text = "☑ क्रय / Purchase      □ इंडेंटर / Indenter"
                    if len(table.rows) > 5 and len(table.rows[5].cells) > 2:
                        _set_cell_text_with_font(table.rows[5].cells[2], processed_text, font_size_pt=11)
                elif case_num == 2:
                    # For Case 2, it's a fixed text indicating processed by purchase dept only
                    processed_text = "केवल क्रयविभागद्वाराप्रक्रिया किया जाएगा / Processed by Purchase Dept Only"
                    if len(table.rows) > 5 and len(table.rows[5].cells) > 2:
                        _set_cell_text_with_font(table.rows[5].cells[2], processed_text, font_size_pt=11)
            except Exception as e:
                print(f"Error setting row 5: {e}")
            
            # Row 6: No Stock Certificate
            try:
                if len(table.rows) > 6 and len(table.rows[6].cells) > 2:
                    _set_cell_text(table.rows[6].cells[2], data.get("no_stock_certificate", ""))
            except Exception as e:
                print(f"Error setting row 6: {e}")
        
        elif case_num == 3:
            # Row 0: Indenting Officer name + Designation
            try:
                if len(table.rows) > 0 and len(table.rows[0].cells) > 2:
                    officer_text = data.get("indenting_officer", "")
                    designation = data.get("designation", "")
                    if designation:
                        officer_text = f"{officer_text}\n{designation}"
                    _set_cell_text(table.rows[0].cells[2], officer_text)
            except Exception as e:
                print(f"Error setting case 3 row 0: {e}")
            
            # Row 1: Centre/Group
            try:
                if len(table.rows) > 1 and len(table.rows[1].cells) > 2:
                    _set_cell_text(table.rows[1].cells[2], data.get("centre_group", ""))
            except Exception as e:
                print(f"Error setting case 3 row 1: {e}")
            
            # Row 2: Material/Service Required + Quantity
            try:
                material_text = material_desc
                if total_quantity:
                    material_text += f"\n{total_quantity}"
                if len(table.rows) > 2 and len(table.rows[2].cells) > 2:
                    _set_cell_text(table.rows[2].cells[2], material_text)
            except Exception as e:
                print(f"Error setting case 3 row 2: {e}")
            
            # Row 3: Project details
            try:
                project_text_parts = []
                if project_number:
                    project_text_parts.append(f"a) {project_number} - {project_name}")
                if data.get("budget_head"):
                    project_text_parts.append(f"b) {data.get('budget_head')}")
                project_text_parts.append(f"c) ₹{cost:,.2f}")
                if len(table.rows) > 3 and len(table.rows[3].cells) > 2:
                    _set_cell_text(table.rows[3].cells[2], "\n".join(project_text_parts))
            except Exception as e:
                print(f"Error setting case 3 row 3: {e}")
            
            # Row 4: Source of supply
            try:
                if len(table.rows) > 4 and len(table.rows[4].cells) > 2:
                    _set_cell_text(table.rows[4].cells[2], data.get("source_of_supply", ""))
            except Exception as e:
                print(f"Error setting case 3 row 4: {e}")
            
            # Row 5: No Stock Certificate
            try:
                if len(table.rows) > 5 and len(table.rows[5].cells) > 2:
                    _set_cell_text(table.rows[5].cells[2], data.get("no_stock_certificate", ""))
            except Exception as e:
                print(f"Error setting case 3 row 5: {e}")
            
            # Row 6: Whether requirement is
            try:
                req_type = data.get("requirement_type", "fresh")
                if req_type == "fresh":
                    req_text = "☑ Fresh     □ Additional   □ Replacement"
                elif req_type == "additional":
                    req_text = "□ Fresh     ☑ Additional   □ Replacement"
                elif req_type == "replacement":
                    req_text = "□ Fresh     □ Additional   ☑ Replacement"
                else:
                    req_text = "☑ Fresh     □ Additional   □ Replacement"
                if len(table.rows) > 6 and len(table.rows[6].cells) > 2:
                    _set_cell_text_with_font(table.rows[6].cells[2], req_text, font_size_pt=11)
            except Exception as e:
                print(f"Error setting case 3 row 6: {e}")
            
            # Row 7: Whether requirement is imported
            try:
                is_imp = data.get("is_imported", "no")
                if is_imp == "yes":
                    imp_text = "☑ Yes        □ No"
                elif is_imp == "no":
                    imp_text = "□ Yes        ☑ No"
                else:
                    imp_text = "□ Yes        ☑ No"
                if len(table.rows) > 7 and len(table.rows[7].cells) > 2:
                    _set_cell_text_with_font(table.rows[7].cells[2], imp_text, font_size_pt=11)
            except Exception as e:
                print(f"Error setting case 3 row 7: {e}")
            
            # Row 8: Delivery period required by
            try:
                if len(table.rows) > 8 and len(table.rows[8].cells) > 2:
                    _set_cell_text_with_font(table.rows[8].cells[2], data.get("delivery_period", ""), font_size_pt=11)
            except Exception as e:
                print(f"Error setting case 3 row 8: {e}")
            
            # Row 9: Proprietary Article Certificate
            try:
                prop_cert = data.get("proprietary_certificate", "not_applicable")
                if prop_cert == "yes":
                    prop_text = "☑ Yes        □ No              □ Not applicable"
                elif prop_cert == "no":
                    prop_text = "□ Yes        ☑ No              □ Not applicable"
                elif prop_cert == "not_applicable":
                    prop_text = "□ Yes        □ No              ☑ Not applicable"
                else:
                    prop_text = "□ Yes        □ No              ☑ Not applicable"
                if len(table.rows) > 9 and len(table.rows[9].cells) > 2:
                    _set_cell_text_with_font(table.rows[9].cells[2], prop_text, font_size_pt=11)
            except Exception as e:
                print(f"Error setting case 3 row 9: {e}")
            
            # Row 10: CGC Approval
            try:
                cgc_app = data.get("cgc_approval", "not_applicable")
                if cgc_app == "yes":
                    cgc_text = "☑ Yes        □ No              □ Not applicable"
                elif cgc_app == "no":
                    cgc_text = "□ Yes        ☑ No              □ Not applicable"
                elif cgc_app == "not_applicable":
                    cgc_text = "□ Yes        □ No              ☑ Not applicable"
                else:
                    cgc_text = "□ Yes        □ No              ☑ Not applicable"
                if len(table.rows) > 10 and len(table.rows[10].cells) > 2:
                    _set_cell_text_with_font(table.rows[10].cells[2], cgc_text, font_size_pt=11)
            except Exception as e:
                print(f"Error setting case 3 row 10: {e}")
    
    # Add Annexure 1 for grouped orders
    if is_grouped and grouped_stocks:
        # Add a page break before annexure
        doc.add_page_break()
        
        # Add Annexure heading
        annexure_heading = doc.add_paragraph()
        annexure_run = annexure_heading.add_run("Annexure 1: Detailed Material Requirements")
        annexure_run.bold = True
        annexure_run.font.size = Pt(14)
        annexure_heading.alignment = 1  # Center
        
        # Create table for annexure
        annexure_table = doc.add_table(rows=1, cols=8)
        annexure_table.style = 'Table Grid'
        
        # Set header row
        headers = ["S.No", "Project Name", "Order Number", "Material", "Process Type", "Form Type", "Dimensions", "Quantity"]
        header_cells = annexure_table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for para in header_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        
        # Add data rows (max 5 per page)
        records_per_page = 5
        total_records = len(grouped_stocks)
        total_pages = (total_records + records_per_page - 1) // records_per_page
        print(f"Annexure pagination: total_records={total_records}, records_per_page={records_per_page}, total_pages={total_pages}")
        
        for page_num in range(total_pages):
            start_idx = page_num * records_per_page
            end_idx = min(start_idx + records_per_page, total_records)
            page_stocks = grouped_stocks[start_idx:end_idx]
            
            # If not first page, add page break
            if page_num > 0:
                doc.add_page_break()
                # Add continuation heading
                cont_heading = doc.add_paragraph()
                cont_run = cont_heading.add_run(f"Annexure 1 (Continued) - Page {page_num + 1}")
                cont_run.bold = True
                cont_run.font.size = Pt(12)
                cont_heading.alignment = 1
                
                # Create new table for this page
                annexure_table = doc.add_table(rows=1, cols=8)
                annexure_table.style = 'Table Grid'
                header_cells = annexure_table.rows[0].cells
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    for para in header_cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
            
            # Add data rows for this page
            for idx, grouped_stock in enumerate(page_stocks):
                row_idx = idx + 1
                if row_idx >= len(annexure_table.rows):
                    annexure_table.add_row()
                
                row_cells = annexure_table.rows[row_idx].cells
                
                # S.No
                row_cells[0].text = str(start_idx + idx + 1)
                
                # Project Name
                project_name = ""
                try:
                    if grouped_stock.source_order_id:
                        order = db.query(OrderModel).filter(OrderModel.id == grouped_stock.source_order_id).first()
                        if order and order.product_id:
                            from DB.models.oms import Product as ProductModel
                            product = db.query(ProductModel).filter(ProductModel.id == order.product_id).first()
                            if product:
                                project_name = product.product_name
                except Exception as e:
                    print(f"Error fetching project name: {e}")
                row_cells[1].text = project_name
                
                # Order Number
                order_number = ""
                try:
                    if grouped_stock.source_order_id:
                        order = db.query(OrderModel).filter(OrderModel.id == grouped_stock.source_order_id).first()
                        if order:
                            order_number = order.sale_order_number or ""
                except Exception as e:
                    print(f"Error fetching order number: {e}")
                row_cells[2].text = order_number
                
                # Material
                material_name = ""
                try:
                    if grouped_stock.material:
                        material_name = grouped_stock.material.material_name or ""
                except Exception as e:
                    print(f"Error fetching material name: {e}")
                row_cells[3].text = material_name
                
                # Process Type
                process_type = grouped_stock.process_type or ""
                row_cells[4].text = process_type
                
                # Form Type
                form_type = grouped_stock.form_type or ""
                row_cells[5].text = form_type
                
                # Dimensions
                dimensions = ""
                if grouped_stock.form_type:
                    if grouped_stock.form_type == 'Round':
                        dimensions = f'Ø{grouped_stock.diameter} × {grouped_stock.length}mm'
                    elif grouped_stock.form_type == 'Square':
                        dimensions = f'{grouped_stock.breadth} × {grouped_stock.height} × {grouped_stock.length}mm'
                    elif grouped_stock.form_type == 'Pipe':
                        dimensions = f'Ø{grouped_stock.outer_diameter}/{grouped_stock.inner_diameter} × {grouped_stock.length}mm'
                row_cells[6].text = dimensions
                
                # Quantity
                row_cells[7].text = str(grouped_stock.quantity) if grouped_stock.quantity else ""
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    filename = f"PurchaseReq_OrderMaterial_{stock_id}_{created_date}.docx"
    
    return buffer, filename
