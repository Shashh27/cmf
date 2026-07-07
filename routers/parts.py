from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from fastapi.responses import Response
from sqlalchemy.orm import Session
from typing import List, Optional
from sqlalchemy.exc import IntegrityError
from sqlalchemy import func, text
from pydantic import BaseModel
import re, tempfile, os, io

from DB.database import get_db, MINIO_BUCKET_NAME
from DB.minio_client import get_minio_client
from DB.models.oms import (
    Part as PartModel, 
    PartType, 
    Order, 
    OrderPartPriority,
    Operation as OperationModel,
    Document as DocumentModel,
    ToolWithPart as ToolWithPartModel,
    OperationDocument as OperationDocumentModel,
    OutSourcePartStatus as OutSourcePartStatusModel,
    Assembly as AssemblyModel,
    DocumentExtractedData as DocumentExtractedDataModel,
)
from DB.models.configuration import PokayokeCompletedLog
from DB.models.inventory import RawMaterial, RawMaterialStock, RawMaterialUnit, Vendors
from DB.models.access_control import AccessUser
from DB.schemas.oms import Part, PartCreate, PartUpdate
from services.stock_auto_update import StockAutoUpdateService
from services.notification_service import NotificationService
from services.raw_material_history_service import RawMaterialHistoryService

router = APIRouter(
    prefix="/parts",
    tags=["parts"]
)


def _check_assembly_recycle_bin_recursive(assembly_id: int, db: Session) -> Optional[str]:
    """
    Check if the assembly or any of its parent assemblies are in the recycle bin.
    Returns the name of the first assembly found in recycle bin, or None if none are.
    """
    current_assembly = db.query(AssemblyModel).filter(AssemblyModel.id == assembly_id).first()
    if not current_assembly:
        return None
    
    # Check current assembly
    if current_assembly.recycle_bin:
        return current_assembly.assembly_name
    
    # Recursively check parent
    if current_assembly.parent_id:
        return _check_assembly_recycle_bin_recursive(current_assembly.parent_id, db)
    
    return None


def _build_part_maps(db: Session):
    """Fetch PartType, RawMaterial, RawMaterialUnit, AccessUser, and Vendors rows once and return id→value maps."""
    type_map = {pt.id: pt.type_name for pt in db.query(PartType).all()}
    rm_map = {rm.id: rm.material_name for rm in db.query(RawMaterial).all()}
    unit_map = {unit.id: unit for unit in db.query(RawMaterialUnit).all()}
    user_map = {u.id: u.user_name for u in db.query(AccessUser).all()}
    vendor_map = {v.id: v.company_name for v in db.query(Vendors).all()}
    return type_map, rm_map, unit_map, user_map, vendor_map


def _part_to_dict(part: PartModel, type_map: dict, rm_map: dict, unit_map: dict, user_map: dict, vendor_map: dict) -> dict:
    # Get unit details if unit is assigned
    unit_details = None
    unit_id = getattr(part, 'raw_material_unit_id', None)
    if unit_id and unit_id in unit_map:
        unit = unit_map[unit_id]
        unit_details = {
            "id": unit.id,
            "total_length": unit.total_length,
            "remaining_length": unit.remaining_length,
            "status": unit.status
        }
    
    return {
        "id": part.id,
        "part_name": part.part_name,
        "part_number": part.part_number,
        "type_id": part.type_id,
        "raw_material_id": part.raw_material_id,
        "raw_material_unit_id": getattr(part, 'raw_material_unit_id', None),
        "required_length": part.required_length,
        "part_detail": part.part_detail,
        "assembly_id": part.assembly_id,
        "product_id": part.product_id,
        "user_id": part.user_id,
        "qty": part.qty,
        "size": part.size,
        "vendor_id": part.vendor_id,
        "type_name": type_map.get(part.type_id),
        "raw_material_name": rm_map.get(part.raw_material_id),
        "raw_material_unit_details": unit_details,
        "user_name": user_map.get(part.user_id) if part.user_id else None,
        "vendor_name": vendor_map.get(part.vendor_id) if part.vendor_id else None,
        "created_at": part.created_at,
        "updated_at": part.updated_at,
    }


@router.post("/", response_model=Part, status_code=status.HTTP_201_CREATED)
def create_part(part: PartCreate, db: Session = Depends(get_db)):
    """Create a new part"""
    # Check if assembly or any parent assembly is in recycle bin - only for assembly parts
    if part.assembly_id:
        recycle_bin_assembly = _check_assembly_recycle_bin_recursive(part.assembly_id, db)
        if recycle_bin_assembly:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Sorry, parts cannot be added because the assembly '{recycle_bin_assembly}' (or a parent assembly) is in the recycle bin. To add parts, please restore the assembly from the recycle bin first."
            )

    # Check for duplicate part number only within the same product
    if part.product_id:
        db_part = db.query(PartModel).filter(
            PartModel.part_number == part.part_number,
            PartModel.product_id == part.product_id
        ).first()
        if db_part:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Part with number {part.part_number} already exists in this product"
            )
    else:
        # If no product_id is provided, check globally (backward compatibility)
        db_part = db.query(PartModel).filter(PartModel.part_number == part.part_number).first()
        if db_part:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Part with number {part.part_number} already exists"
            )

    # Check if raw material allocation is needed
    needs_allocation = getattr(part, 'raw_material_unit_id', None) and getattr(part, 'required_length', None)
    
    if needs_allocation:
        # Validate unit availability before creating part
        unit = db.query(RawMaterialUnit).filter(RawMaterialUnit.id == part.raw_material_unit_id).first()
        if not unit:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Unit with id {part.raw_material_unit_id} not found"
            )
        
        if unit.remaining_length < getattr(part, 'required_length', 0):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Not enough length in unit. Required: {getattr(part, 'required_length', 0)}, Available: {unit.remaining_length}"
            )

    # Create the part first
    db_part = PartModel(**part.model_dump())
    db.add(db_part)
    db.commit()
    db.refresh(db_part)
    
    # Now do the allocation since we have the part ID
    if needs_allocation:
        # Unit-based allocation is handled by the assign-material endpoint
        # No automatic allocation needed here
        pass

    # =================================================================
    # IMPORTANT: OrderPartPriority auto-creation is PERMANENTLY DISABLED
    # Parts should NOT automatically create entries in order_part_priorities table
    # This table should only be managed manually through the priority management UI
    # =================================================================
    # Automatic OrderPartPriority creation disabled
    # if db_part.product_id and db_part.type_id:
    #     part_type = db.query(PartType).filter(PartType.id == db_part.type_id).first()
    #     if part_type and part_type.type_name and part_type.type_name.lower() == "in-house":
    #         orders = db.query(Order).filter(Order.product_id == db_part.product_id).all()
    #         max_priority = db.query(func.max(OrderPartPriority.priority)).scalar() or 0
    #         for index, order in enumerate(orders):
    #             priority_entry = OrderPartPriority(
    #                 order_id=order.id,
    #                 product_id=db_part.product_id,
    #                 part_id=db_part.id,
    #                 priority=max_priority + 1 + index,
    #             )
    #             db.add(priority_entry)
    #         db.commit()

    # EXTRA SAFETY: Remove any OrderPartPriority entries that might have been created
    # for this part by triggers or other mechanisms
    db.query(OrderPartPriority).filter(OrderPartPriority.part_id == db_part.id).delete(
        synchronize_session=False
    )
    db.commit()

    # Log part creation for PC notifications
    user_name = None
    user_role = None
    if db_part.user_id:
        user = db.query(AccessUser).filter(AccessUser.id == db_part.user_id).first()
        user_name = user.user_name if user else None
        user_role = user.role if user else None
    
    NotificationService.log_part_change(
        db=db,
        part_id=db_part.id,
        action="created",
        user_id=db_part.user_id,
        user_name=user_name,
        user_role=user_role,
        details={"part_name": db_part.part_name, "part_number": db_part.part_number}
    )

    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return _part_to_dict(db_part, type_map, rm_map, unit_map, user_map, vendor_map)


@router.get("/", response_model=List[Part])
def get_parts(user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts with type, raw material, and user names. Filter by user_id for module-specific views."""
    query = db.query(PartModel).order_by(PartModel.id.asc())
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, unit_map, user_map, vendor_map) for p in parts]


@router.get("/{part_id}", response_model=Part)
def get_part(part_id: int, db: Session = Depends(get_db)):
    """Get a specific part by ID with type, raw_material, and user names."""
    part = db.query(PartModel).filter(PartModel.id == part_id).first()
    if not part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with id {part_id} not found"
        )
    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return _part_to_dict(part, type_map, rm_map, unit_map, user_map, vendor_map)


@router.get("/product/{product_id}", response_model=List[Part])
def get_parts_by_product(product_id: int, user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts for a specific product. Filter by user_id for module-specific views."""
    query = db.query(PartModel).filter(PartModel.product_id == product_id)
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, unit_map, user_map, vendor_map) for p in parts]


@router.get("/assembly/{assembly_id}", response_model=List[Part])
def get_parts_by_assembly(assembly_id: int, user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts for a specific assembly. Filter by user_id for module-specific views."""
    query = db.query(PartModel).filter(PartModel.assembly_id == assembly_id)
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, unit_map, user_map, vendor_map) for p in parts]


@router.get("/type/{type_id}", response_model=List[Part])
def get_parts_by_type(type_id: int, user_id: int | None = None, db: Session = Depends(get_db)):
    """Get all parts of a specific type. Filter by user_id for module-specific views."""
    query = db.query(PartModel).filter(PartModel.type_id == type_id)
    if user_id is not None:
        query = query.filter(PartModel.user_id == user_id)
    parts = query.all()
    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return [_part_to_dict(p, type_map, rm_map, unit_map, user_map, vendor_map) for p in parts]


@router.put("/{part_id}", response_model=Part)
def update_part(part_id: int, part: PartUpdate, db: Session = Depends(get_db)):
    """Update a part"""
    db_part = db.query(PartModel).filter(PartModel.id == part_id).first()
    if not db_part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with id {part_id} not found"
        )

    # Check if part is scheduled (status is "active")
    schedule_status = db.execute(
        text("SELECT status FROM scheduling.part_schedule_status WHERE part_id = :pid"),
        {"pid": part_id}
    ).fetchone()
    
    if schedule_status and schedule_status[0] and schedule_status[0].lower() == "active":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Sorry, this part cannot be modified because it is currently scheduled for production. To make changes, please inactivate the part's schedule status first."
        )

    update_data = part.model_dump(exclude_unset=True)
    
    # Capture old values before updating
    old_values = {}
    for field in update_data.keys():
        old_values[field] = getattr(db_part, field, None)
    
    # Check if we're switching from outsource to in-house
    is_switching_to_inhouse = (
        'part_detail' in update_data and
        'type_id' in update_data
    )
    
    if is_switching_to_inhouse:
        # Check if old part was outsource and new part is in-house
        new_type_id = update_data.get('type_id')
        old_part = db_part
        
        # Query PartType table to get type names
        from DB.models.oms import PartType
        old_type = db.query(PartType).filter(PartType.id == old_part.type_id).first()
        new_type = db.query(PartType).filter(PartType.id == new_type_id).first()
        
        old_is_outsource = 'out' in old_type.type_name.lower() if old_type else False
        new_is_inhouse = 'in-house' in new_type.type_name.lower() if new_type else False
        
        if old_is_outsource and new_is_inhouse:
            # Clear vendor_id when switching from outsource to in-house
            update_data['vendor_id'] = None
            print("Clearing vendor_id when switching from outsource to in-house")
    
    # Check if we're switching to WITHOUT_RAW_MATERIAL by checking part_detail
    is_switching_to_without_raw = (
        'part_detail' in update_data and 
        update_data.get('part_detail') == 'WITHOUT_RAW_MATERIAL'
    )
    
    # Check if raw material allocation is being updated
    is_updating_raw_material = any(key in update_data for key in [
        'raw_material_unit_id', 'raw_material_id', 'qty'
    ])
    
    # Check if all raw material fields are being cleared (set to null)
    is_clearing_raw_material = (
        update_data.get('raw_material_id') is None and
        update_data.get('raw_material_unit_id') is None and
        update_data.get('required_length') is None
    )
    
    # Also consider clearing if switching to WITHOUT_RAW_MATERIAL - this takes priority
    if is_switching_to_without_raw:
        is_clearing_raw_material = True
    
    # If updating raw material allocation, handle the allocation logic
    unit = None  # Initialize unit variable
    if is_updating_raw_material or is_clearing_raw_material:
        # Special case: All raw material fields are being cleared
        if is_clearing_raw_material:
            # If there's an existing unit assignment, restore it
            material_name_for_history = None
            unit_id_for_history = None
            if db_part.raw_material_unit_id:
                # Get the unit
                unit = db.query(RawMaterialUnit).filter(RawMaterialUnit.id == db_part.raw_material_unit_id).first()
                if unit:
                    # Store for history before clearing
                    unit_id_for_history = unit.id
                    material_name_for_history = unit.stock.material.material_name if unit.stock and unit.stock.material else "Unknown"
                    
                    # Get usage record
                    from DB.models.inventory import RawMaterialUsage as RawMaterialUsageModel
                    usage = db.query(RawMaterialUsageModel).filter(
                        RawMaterialUsageModel.raw_material_unit_id == db_part.raw_material_unit_id,
                        RawMaterialUsageModel.part_id == part_id
                    ).first()

                    if usage:
                        # Restore unit's remaining length
                        unit.remaining_length += usage.used_length

                        # Update unit status
                        if unit.remaining_length == unit.total_length:
                            unit.status = "available"
                        elif unit.remaining_length > 0:
                            unit.status = "partially_used"

                        # Delete usage record
                        db.delete(usage)

            # Clear part fields
            db_part.raw_material_id = None
            db_part.raw_material_unit_id = None
            db_part.required_length = None
            
            # Log history for material unlinking
            if unit_id_for_history:
                try:
                    RawMaterialHistoryService.log_material_unlinked(
                        db=db,
                        unit_id=unit_id_for_history,
                        part_id=part_id,
                        material_name=material_name_for_history,
                        user_id=db_part.user_id
                    )
                except Exception as e:
                    print(f"Error logging material unlinking history: {e}")
    
    # Update other fields normally
    for field, value in update_data.items():
        setattr(db_part, field, value)

    db.commit()
    
    # Log part update for PC notifications
    user_name = None
    user_role = None
    if db_part.user_id:
        user = db.query(AccessUser).filter(AccessUser.id == db_part.user_id).first()
        user_name = user.user_name if user else None
        user_role = user.role if user else None
    
    # Capture changes with old and new values
    changes = {}
    for field in update_data.keys():
        old_value = old_values[field]
        new_value = update_data[field]
        
        # Convert time/datetime objects to strings for JSON serialization
        if old_value is not None and hasattr(old_value, 'isoformat'):
            old_value = old_value.isoformat()
        elif hasattr(old_value, '__str__') and not isinstance(old_value, (str, int, float, bool)):
            old_value = str(old_value)
        
        if new_value is not None and hasattr(new_value, 'isoformat'):
            new_value = new_value.isoformat()
        elif hasattr(new_value, '__str__') and not isinstance(new_value, (str, int, float, bool)):
            new_value = str(new_value)
        
        changes[field] = {"old": old_value, "new": new_value}
    
    NotificationService.log_part_change(
        db=db,
        part_id=db_part.id,
        action="updated",
        user_id=db_part.user_id,
        user_name=user_name,
        user_role=user_role,
        details={"part_name": db_part.part_name, "part_number": db_part.part_number, "changes": changes}
    )
    
    # 🔥 Update stock status based on unit statuses if raw material was cleared
    if is_clearing_raw_material and unit:
        StockAutoUpdateService.update_stock_status_from_units(db, unit.stock_id)
    db.refresh(db_part)
    type_map, rm_map, unit_map, user_map, vendor_map = _build_part_maps(db)
    return _part_to_dict(db_part, type_map, rm_map, unit_map, user_map, vendor_map)


@router.delete("/{part_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_part(part_id: int, db: Session = Depends(get_db)):
    """Delete a part and all its references across all schemas with raw material stock restoration."""
    db_part = db.query(PartModel).filter(PartModel.id == part_id).first()
    if not db_part:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Part with id {part_id} not found"
        )

    # Log part deletion for PC notifications before deletion
    user_name = None
    user_role = None
    if db_part.user_id:
        user = db.query(AccessUser).filter(AccessUser.id == db_part.user_id).first()
        user_name = user.user_name if user else None
        user_role = user.role if user else None
    
    NotificationService.log_part_change(
        db=db,
        part_id=db_part.id,
        action="deleted",
        user_id=db_part.user_id,
        user_name=user_name,
        user_role=user_role,
        details={"part_name": db_part.part_name, "part_number": db_part.part_number}
    )

    try:
        # 1. Delete pokayoke logs for this part
        result = db.execute(
            text(
                "SELECT id FROM configuration.pokayoke_completed_logs "
                "WHERE part_id = :pid"
            ),
            {"pid": part_id},
        )
        log_ids = [row[0] for row in result]
        for log_id in log_ids:
            log_obj = (
                db.query(PokayokeCompletedLog)
                .filter(PokayokeCompletedLog.id == log_id)
                .first()
            )
            if log_obj:
                db.delete(log_obj)
        db.flush()

        # 2. Delete from scheduling.part_schedule_status
        db.execute(
            text("DELETE FROM scheduling.part_schedule_status WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 3. Delete order part priorities
        db.query(OrderPartPriority).filter(OrderPartPriority.part_id == part_id).delete(
            synchronize_session=False
        )

        # 4. Delete operations and their documents/tools
        operations = db.query(OperationModel).filter(OperationModel.part_id == part_id).all()
        operation_ids = [op.id for op in operations]
        if operation_ids:
            # Delete from production_monitoring.machine_live_history
            db.execute(
                text("DELETE FROM production_monitoring.machine_live_history WHERE current_operation_id IN :op_ids"),
                {"op_ids": tuple(operation_ids)}
            )

            # Delete from production_monitoring.machine_live_status
            db.execute(
                text("DELETE FROM production_monitoring.machine_live_status WHERE current_operation_id IN :op_ids"),
                {"op_ids": tuple(operation_ids)}
            )

            # Delete from scheduling.production_logs
            db.execute(
                text("DELETE FROM scheduling.production_logs WHERE operation_id IN :op_ids"),
                {"op_ids": tuple(operation_ids)}
            )

            # Delete from scheduling.rescheduling_items (by operation_id)
            db.execute(
                text("DELETE FROM scheduling.rescheduling_items WHERE operation_id IN :op_ids"),
                {"op_ids": tuple(operation_ids)}
            )

            # Delete from scheduling.planned_schedule_items
            db.execute(
                text("DELETE FROM scheduling.planned_schedule_items WHERE operation_id IN :op_ids"),
                {"op_ids": tuple(operation_ids)}
            )

            db.query(OperationDocumentModel).filter(
                OperationDocumentModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            
            db.query(ToolWithPartModel).filter(
                ToolWithPartModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            
            db.query(OperationModel).filter(OperationModel.id.in_(operation_ids)).delete(
                synchronize_session=False
            )

        # 5. Delete part documents and their extracted data
        db.query(DocumentExtractedDataModel).filter(
            DocumentExtractedDataModel.part_id == part_id
        ).delete(synchronize_session=False)

        db.query(DocumentModel).filter(DocumentModel.part_id == part_id).delete(
            synchronize_session=False
        )

        # 6. Delete out source part status records
        db.query(OutSourcePartStatusModel).filter(
            OutSourcePartStatusModel.part_id == part_id
        ).delete(synchronize_session=False)

        # 7. Delete from maintenance.component_issues
        db.execute(
            text("DELETE FROM maintenance.component_issues WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 8. Delete from maintenance.help_support
        db.execute(
            text("DELETE FROM maintenance.help_support WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 9. Delete tools with part (that are not associated with operations)
        db.query(ToolWithPartModel).filter(ToolWithPartModel.part_id == part_id).delete(
            synchronize_session=False
        )

        # 10. Delete inventory requests and return requests
        db.execute(
            text("DELETE FROM inventory.inventory_return_requests WHERE requested_id IN (SELECT id FROM inventory.inventory_requests WHERE part_id = :pid)"),
            {"pid": part_id}
        )
        db.execute(
            text("DELETE FROM inventory.inventory_requests WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 11. Delete from inventory.raw_material_usage
        db.execute(
            text("DELETE FROM inventory.raw_material_usage WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 12. Delete from notifications.inspection_notifications (by part_number)
        if db_part.part_number:
            db.execute(
                text("DELETE FROM notifications.inspection_notifications WHERE part_number = :pnum"),
                {"pnum": db_part.part_number}
            )

        # 13. Delete from quality.inspection_plan_status (by part_number)
        if db_part.part_number:
            db.execute(
                text("DELETE FROM quality.inspection_plan_status WHERE part_number = :pnum"),
                {"pnum": db_part.part_number}
            )

        # 14. Delete from quality.master_boc (part_id is character varying)
        db.execute(
            text("DELETE FROM quality.master_boc WHERE part_id = :pid"),
            {"pid": str(part_id)}
        )

        # 15. Delete from quality.notes
        db.execute(
            text("DELETE FROM quality.notes WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 16. Delete from quality.stage_inspection
        db.execute(
            text("DELETE FROM quality.stage_inspection WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 17. Delete from scheduling.machine_schedule
        db.execute(
            text("DELETE FROM scheduling.machine_schedule WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 18. Delete from scheduling.operation_status
        db.execute(
            text("DELETE FROM scheduling.operation_status WHERE part_id = :pid"),
            {"pid": part_id}
        )

        # 19. Delete from scheduling.rescheduling_items (by part_id and part_number)
        db.execute(
            text("DELETE FROM scheduling.rescheduling_items WHERE part_id = :pid"),
            {"pid": part_id}
        )
        if db_part.part_number:
            db.execute(
                text("DELETE FROM scheduling.rescheduling_items WHERE part_number = :pnum"),
                {"pnum": db_part.part_number}
            )

        # 20. Deallocate raw material if this part has any allocated (RESTORE STOCK)
        stock_id_to_update = None
        if db_part.raw_material_unit_id:
            try:
                unit = db.query(RawMaterialUnit).filter(RawMaterialUnit.id == db_part.raw_material_unit_id).first()
                if unit:
                    stock_id_to_update = unit.stock_id
                    usage = db.query(RawMaterialUsageModel).filter(
                        RawMaterialUsageModel.raw_material_unit_id == db_part.raw_material_unit_id,
                        RawMaterialUsageModel.part_id == part_id
                    ).first()
                    if usage:
                        # RESTORE: Add back the used length to unit's remaining length
                        unit.remaining_length += usage.used_length
                        
                        # Update unit status based on restored length
                        if unit.remaining_length == unit.total_length:
                            unit.status = "available"
                        elif unit.remaining_length > 0:
                            unit.status = "partially_used"
                        
                        # Delete the usage record
                        db.delete(usage)
            except Exception as e:
                # Log error but don't fail the deletion
                print(f"Warning: Could not deallocate raw material: {e}")

        # Finally, delete the part itself
        db.delete(db_part)
        db.commit()
        
        # Update stock status based on unit statuses after part deletion
        if stock_id_to_update:
            StockAutoUpdateService.update_stock_status_from_units(db, stock_id_to_update)
    except Exception as e:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Error deleting part: {str(e)}"
        )
    return None



# ─────────────────────────────────────────────────────────────────────────────
# UPDATED /parse-doc endpoint for the NEW Parts Template
#
# New template columns (6 cols, col index 0-based):
#   0: Sl.No
#   1: Name of Part
#   2: Part No
#   3: Size
#   4: Part Quantity
#   5: Part Type (In-house/Out-source/Standard)
#
# Part Type normalisation handles ALL case variations:
#   "In-house", "INHOUSE", "in house", "In House"  → type_id = 1
#   "Out-source", "OUTSOURCE", "out source"          → type_id = 2
#   "Standard", "STANDARD", "standard"               → type_id = 3
# ─────────────────────────────────────────────────────────────────────────────

# ADD THESE IMPORTS at the top of parts.py (merge with existing imports)
# from fastapi import UploadFile, File
# import tempfile, os, re
# from docx import Document

@router.post("/parse-doc", status_code=status.HTTP_200_OK)
async def parse_parts_doc(file: UploadFile = File(...)):
    """
    Accept a .docx BOM file and return extracted part rows.

    New Template Structure (6 columns):
    ─────────────────────────────────────────────────────
    Col 0  → Sl.No
    Col 1  → Name of Part
    Col 2  → Part No
    Col 3  → Size
    Col 4  → Part Quantity
    Col 5  → Part Type (In-house/Out-source/Standard)

    Row 0 is always the header row.
    Rows with empty Name of Part AND empty Part No are skipped (blank filler rows).
    """
    from docx import Document

    suffix = os.path.splitext(file.filename or "")[-1].lower()
    if suffix != ".docx":
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported. Please convert .doc to .docx first.",
        )

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        doc = Document(tmp_path)
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Could not open document: {exc}",
        )
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    if not doc.tables:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail="No tables found in the document.",
        )

    # ── Helpers ──────────────────────────────────────────────────────────────

    def _cell(row_cells: list, idx: int) -> str:
        """Return cleaned cell text at index, or '' if out of range."""
        if idx < 0 or idx >= len(row_cells):
            return ""
        return re.sub(r"\s+", " ", row_cells[idx]).strip()

    def _is_header(cells: list) -> bool:
        """Detect header row by checking for 'name of part' in joined text."""
        joined = " ".join(c.lower() for c in cells)
        return "name of part" in joined

    def _normalize_part_type(raw: str) -> int:
        """
        Map any casing/spacing variation of part type text to a type_id.

        Accepted values (case-insensitive, ignores hyphens/spaces):
          in-house  / in house  / inhouse   → 1
          out-source / out source / outsource → 2
          standard                            → 3

        Defaults to 1 (In-house) if unrecognised or empty.
        """
        if not raw:
            return 1

        # Normalise: lowercase, strip hyphens and extra spaces
        normalised = re.sub(r"[-\s]+", "", raw.lower())

        if "inhouse" in normalised:
            return 1
        if "outsource" in normalised:
            return 2
        if "standard" in normalised:
            return 3

        return 1  # default

    # ── Column indices for the new template ──────────────────────────────────
    COL_SLNO     = 0
    COL_NAME     = 1
    COL_PARTNO   = 2
    COL_SIZE     = 3
    COL_QTY      = 4
    COL_TYPE     = 5

    parts: list[dict] = []

    for table in doc.tables:
        all_rows = [[c.text for c in row.cells] for row in table.rows]

        # Skip tables that don't look like a parts BOM
        if not any(_is_header(r) for r in all_rows):
            continue

        for row_cells in all_rows:
            # Skip header rows
            if _is_header(row_cells):
                continue

            part_name   = _cell(row_cells, COL_NAME)
            part_number = _cell(row_cells, COL_PARTNO)
            qty_raw     = _cell(row_cells, COL_QTY)
            size_raw    = _cell(row_cells, COL_SIZE)
            type_raw    = _cell(row_cells, COL_TYPE)

            # Skip blank filler rows (both name and number empty)
            if not part_name and not part_number:
                continue

            # Skip rows where the entire row is effectively one value
            # (e.g., serial-number-only rows: only Sl.No is filled)
            non_empty = [c for c in row_cells if c.strip()]
            if len(non_empty) == 1:
                continue

            # Parse quantity: extract the first integer
            # Handles formats like "57", "57 (55+2)", "50\n(49+1)"
            qty: int | None = None
            m = re.search(r"\d+", qty_raw.replace(",", ""))
            if m:
                qty = int(m.group())

            parts.append(
                {
                    "part_name":         part_name,
                    "part_number":       part_number,
                    "qty":               qty,
                    "size":              size_raw or None,
                    "raw_material_name": None,          # not in new template
                    "type_id":           _normalize_part_type(type_raw),
                    "part_detail":       None,
                }
            )

    if not parts:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=(
                "No part rows could be extracted. "
                "Ensure the BOM table has columns: "
                "'Sl.No', 'Name of Part', 'Part No', 'Size', "
                "'Part Quantity', 'Part Type (In-house/Out-source/Standard)'."
            ),
        )

    return {"parts": parts, "count": len(parts)}


# ── Part-type normalisation summary ──────────────────────────────────────────
#
#  Input (any case)              │ Stored type_id │ Display label
# ───────────────────────────────┼────────────────┼──────────────
#  In-house / IN-HOUSE / inhouse │       1        │ In-house
#  Out-source / OUTSOURCE        │       2        │ Out-source
#  Standard / STANDARD           │       3        │ Standard
#  (empty / unrecognised)        │       1        │ In-house  ← default
# ─────────────────────────────────────────────────────────────────────────────




 
# ══════════════════════════════════════════════════════════════════════════════
#  2.  NEW  /bulk  — create many parts in a single POST request
# ══════════════════════════════════════════════════════════════════════════════
class BulkPartCreateItem(BaseModel):
    """One part inside a bulk-create request."""
    part_name:       str
    part_number:     str
    type_id:         int        = 1
    raw_material_id: int | None = None
    part_detail:     str | None = None
    assembly_id:     int | None = None
    product_id:      int | None = None
    user_id:         int | None = None
    qty:             int | None = None
    size:            str | None = None
 
 
class BulkPartCreateRequest(BaseModel):
    parts: list[BulkPartCreateItem]
 
 
class BulkPartCreateResult(BaseModel):
    created:    list[dict]   # successfully created parts
    duplicates: list[str]    # part_numbers that already existed (skipped)
    errors:     list[dict]   # [{"part_number": …, "error": …}]
 
 
@router.post("/bulk", response_model=BulkPartCreateResult, status_code=status.HTTP_200_OK)
def bulk_create_parts(payload: BulkPartCreateRequest, db: Session = Depends(get_db)):
    """
    Create multiple parts in a single database transaction (one HTTP call).
 
    Instead of calling POST /parts/ N times, the frontend sends all rows here
    at once.  Returns three lists so the UI can show per-row status:
      • created    – parts successfully inserted
      • duplicates – part_numbers that already existed (skipped, not a fatal error)
      • errors     – parts that failed for any other reason
    """
    created:    list[dict] = []
    duplicates: list[str]  = []
    errors:     list[dict] = []

    # Check if any assembly or parent assembly is in recycle bin - only for assembly parts
    assembly_ids = set(item.assembly_id for item in payload.parts if item.assembly_id)
    if assembly_ids:
        recycle_bin_assemblies = []
        for assembly_id in assembly_ids:
            recycle_bin_assembly = _check_assembly_recycle_bin_recursive(assembly_id, db)
            if recycle_bin_assembly and recycle_bin_assembly not in recycle_bin_assemblies:
                recycle_bin_assemblies.append(recycle_bin_assembly)
        if recycle_bin_assemblies:
            assembly_names_str = ', '.join(recycle_bin_assemblies)
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Sorry, parts cannot be added because the following assemblies (or their parents) are in the recycle bin: {assembly_names_str}. To add parts, please restore these assemblies from the recycle bin first."
            )

 
    type_map, rm_map, stock_map, user_map, vendor_map = _build_part_maps(db)
 
    # Pre-check all part numbers in ONE query to avoid N round-trips
    # Check for duplicates only within the same product
    incoming_numbers = [item.part_number for item in payload.parts if item.part_number]
    existing_numbers: set[str] = set()
    if incoming_numbers:
        # Group parts by product_id to check duplicates within each product
        product_groups = {}
        for item in payload.parts:
            if item.part_number:
                if item.product_id not in product_groups:
                    product_groups[item.product_id] = []
                product_groups[item.product_id].append(item.part_number)
        
        # Check existing numbers for each product group
        for product_id, numbers in product_groups.items():
            if product_id is not None:
                # Check within specific product
                existing_in_product = {
                    row.part_number
                    for row in db.query(PartModel.part_number)
                    .filter(
                        PartModel.part_number.in_(numbers),
                        PartModel.product_id == product_id
                    )
                    .all()
                }
                existing_numbers.update(existing_in_product)
            else:
                # If no product_id, check globally (backward compatibility)
                existing_global = {
                    row.part_number
                    for row in db.query(PartModel.part_number)
                    .filter(PartModel.part_number.in_(numbers))
                    .all()
                }
                existing_numbers.update(existing_global)

    for item in payload.parts:
        # Intra-batch duplicate check (catches repeated numbers within the same payload and product)
        if item.part_number in existing_numbers:
            duplicates.append(item.part_number)
            continue
 
        try:
            db_part = PartModel(**item.model_dump())
            db.add(db_part)
            db.flush()   # assign id without committing yet
            created.append(_part_to_dict(db_part, type_map, rm_map, stock_map, user_map, vendor_map))
            existing_numbers.add(item.part_number)   # prevent intra-batch dupes
        except Exception as exc:
            db.rollback()
            errors.append({"part_number": item.part_number, "error": str(exc)})
            type_map, rm_map, stock_map, user_map, vendor_map = _build_part_maps(db)
 
    # Single commit for all successful inserts
    if created:
        try:
            db.commit()
        except Exception as exc:
            db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Commit failed: {exc}",
            )
 
    return BulkPartCreateResult(created=created, duplicates=duplicates, errors=errors)


class BulkDeleteResult(BaseModel):
    assembly_id: int | None = None
    product_id: int | None = None
    deleted_count: int
    part_ids: list[int]


@router.delete(
    "/bulk-by-assembly/{assembly_id}",
    response_model=BulkDeleteResult,
    status_code=status.HTTP_200_OK,
)
def bulk_delete_parts_by_assembly(assembly_id: int, db: Session = Depends(get_db)):
    """
    Delete ALL parts linked to the given assembly_id in a single call,
    along with every dependent record for each part — identical cleanup
    logic to the single-part DELETE /{part_id} endpoint.

    Called from the Assembly Document Panel when the user clicks
    "Delete All Parts" for an assembly.

    Returns the count and IDs of deleted parts so the frontend can
    update its local state without a refetch.
    """
    # ── 1. Find all part IDs for this assembly ────────────────────────────────
    parts = (
        db.query(PartModel)
        .filter(PartModel.assembly_id == assembly_id)
        .all()
    )

    if not parts:
        # Nothing to delete — return gracefully (not a 404)
        return BulkDeleteResult(
            assembly_id=assembly_id,
            deleted_count=0,
            part_ids=[],
        )

    part_ids = [p.id for p in parts]

    try:
        # ── 2. Pokayoke logs ──────────────────────────────────────────────────
        for part_id in part_ids:
            result = db.execute(
                text(
                    "SELECT id FROM configuration.pokayoke_completed_logs "
                    "WHERE part_id = :pid"
                ),
                {"pid": part_id},
            )
            log_ids = [row[0] for row in result]
            for log_id in log_ids:
                log_obj = (
                    db.query(PokayokeCompletedLog)
                    .filter(PokayokeCompletedLog.id == log_id)
                    .first()
                )
                if log_obj:
                    db.delete(log_obj)
        db.flush()

        # ── 3. Scheduling: part_schedule_status ──────────────────────────────
        db.execute(
            text(
                "DELETE FROM scheduling.part_schedule_status "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 4. Order part priorities ──────────────────────────────────────────
        db.query(OrderPartPriority).filter(
            OrderPartPriority.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 5. Operations → planned_schedule_items, documents, tools ─────────
        operations = (
            db.query(OperationModel)
            .filter(OperationModel.part_id.in_(part_ids))
            .all()
        )
        operation_ids = [op.id for op in operations]

        if operation_ids:
            db.execute(
                text(
                    "DELETE FROM scheduling.planned_schedule_items "
                    "WHERE operation_id = ANY(:oids)"
                ),
                {"oids": operation_ids},
            )
            db.query(OperationDocumentModel).filter(
                OperationDocumentModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            db.query(ToolWithPartModel).filter(
                ToolWithPartModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            db.query(OperationModel).filter(
                OperationModel.id.in_(operation_ids)
            ).delete(synchronize_session=False)

        # ── 6. Part documents + extracted data ───────────────────────────────
        db.query(DocumentExtractedDataModel).filter(
            DocumentExtractedDataModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)
        db.query(DocumentModel).filter(
            DocumentModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 7. Out-source part statuses ─────────────────────────────────────
        db.query(OutSourcePartStatusModel).filter(
            OutSourcePartStatusModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 8. Maintenance component issues ──────────────────────────────────
        db.execute(
            text(
                "DELETE FROM maintenance.component_issues "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 10. Tools with part (not operation-linked) ────────────────────────
        db.query(ToolWithPartModel).filter(
            ToolWithPartModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 11. Inventory return requests + inventory requests ────────────────
        db.execute(
            text(
                "DELETE FROM inventory.inventory_return_requests "
                "WHERE requested_id IN ("
                "  SELECT id FROM inventory.inventory_requests "
                "  WHERE part_id = ANY(:pids)"
                ")"
            ),
            {"pids": part_ids},
        )
        db.execute(
            text(
                "DELETE FROM inventory.inventory_requests "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 12. Finally delete all parts ──────────────────────────────────────
        db.query(PartModel).filter(
            PartModel.id.in_(part_ids)
        ).delete(synchronize_session=False)

        db.commit()

    except Exception as exc:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Bulk delete failed: {str(exc)}",
        )

    return BulkDeleteResult(
        assembly_id=assembly_id,
        deleted_count=len(part_ids),
        part_ids=part_ids,
    )


@router.delete(
    "/bulk-by-product/{product_id}",
    response_model=BulkDeleteResult,
    status_code=status.HTTP_200_OK,
)
def bulk_delete_parts_by_product(product_id: int, db: Session = Depends(get_db)):
    """
    Delete ALL parts linked to given product_id in a single call,
    along with every dependent record for each part — identical cleanup
    logic to single-part DELETE /{part_id} endpoint.

    Called from BillOfMaterials when user clicks
    "Delete All Parts" for a product.

    Returns count and IDs of deleted parts so the frontend can
    update its local state without a refetch.
    """
    # ── 1. Find all part IDs for this product ────────────────────────────────
    parts = (
        db.query(PartModel)
        .filter(PartModel.product_id == product_id)
        .all()
    )

    if not parts:
        # Nothing to delete — return gracefully (not a 404)
        return BulkDeleteResult(
            product_id=product_id,
            deleted_count=0,
            part_ids=[],
        )

    part_ids = [p.id for p in parts]

    try:
        # ── 2. Pokayoke logs ──────────────────────────────────────────────────
        for part_id in part_ids:
            result = db.execute(
                text(
                    "SELECT id FROM configuration.pokayoke_completed_logs "
                    "WHERE part_id = :pid"
                ),
                {"pid": part_id},
            )
            log_ids = [row[0] for row in result]
            for log_id in log_ids:
                log_obj = (
                    db.query(PokayokeCompletedLog)
                    .filter(PokayokeCompletedLog.id == log_id)
                    .first()
                )
                if log_obj:
                    db.delete(log_obj)
        db.flush()

        # ── 3. Scheduling: part_schedule_status ──────────────────────────────
        db.execute(
            text(
                "DELETE FROM scheduling.part_schedule_status "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 4. Order part priorities ──────────────────────────────────────────
        db.query(OrderPartPriority).filter(
            OrderPartPriority.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 5. Operations → planned_schedule_items, documents, tools ─────────
        operations = (
            db.query(OperationModel)
            .filter(OperationModel.part_id.in_(part_ids))
            .all()
        )
        operation_ids = [op.id for op in operations]

        if operation_ids:
            db.execute(
                text(
                    "DELETE FROM scheduling.planned_schedule_items "
                    "WHERE operation_id = ANY(:oids)"
                ),
                {"oids": operation_ids},
            )
            db.query(OperationDocumentModel).filter(
                OperationDocumentModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            db.query(ToolWithPartModel).filter(
                ToolWithPartModel.operation_id.in_(operation_ids)
            ).delete(synchronize_session=False)
            db.query(OperationModel).filter(
                OperationModel.id.in_(operation_ids)
            ).delete(synchronize_session=False)

        # ── 6. Part documents + extracted data ───────────────────────────────
        db.query(DocumentExtractedDataModel).filter(
            DocumentExtractedDataModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)
        db.query(DocumentModel).filter(
            DocumentModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 7. Out-source part statuses ─────────────────────────────────────
        db.query(OutSourcePartStatusModel).filter(
            OutSourcePartStatusModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 8. Maintenance component issues ──────────────────────────────────
        db.execute(
            text(
                "DELETE FROM maintenance.component_issues "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 10. Tools with part (not operation-linked) ────────────────────────
        db.query(ToolWithPartModel).filter(
            ToolWithPartModel.part_id.in_(part_ids)
        ).delete(synchronize_session=False)

        # ── 11. Inventory return requests + inventory requests ────────────────
        db.execute(
            text(
                "DELETE FROM inventory.inventory_return_requests "
                "WHERE requested_id IN ("
                "  SELECT id FROM inventory.inventory_requests "
                "  WHERE part_id = ANY(:pids)"
                ")"
            ),
            {"pids": part_ids},
        )
        db.execute(
            text(
                "DELETE FROM inventory.inventory_requests "
                "WHERE part_id = ANY(:pids)"
            ),
            {"pids": part_ids},
        )

        # ── 12. Finally delete all parts ──────────────────────────────────────
        db.query(PartModel).filter(
            PartModel.id.in_(part_ids)
        ).delete(synchronize_session=False)

        db.commit()

    except Exception as exc:
        db.rollback()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Bulk delete failed: {str(exc)}",
        )

    return BulkDeleteResult(
        product_id=product_id,
        deleted_count=len(part_ids),
        part_ids=part_ids,
    )


# ── Parts Template ───────────────────────────────────────────────────────────

PARTS_TEMPLATE_OBJECT_NAME = "templates/parts_template.docx"


@router.post("/template/upload")
async def upload_parts_template(file: UploadFile = File(...)):
    """
    Upload the parts template file to MinIO.
    This replaces any existing template file.
    """
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are allowed for the parts template"
        )

    content = await file.read()
    if not content:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Empty file"
        )

    minio_client = get_minio_client()

    try:
        # Upload to MinIO with fixed object name
        file_stream = io.BytesIO(content)
        url = minio_client.upload_file(
            file_data=file_stream,
            object_name=PARTS_TEMPLATE_OBJECT_NAME,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        return {
            "message": "Parts template uploaded successfully",
            "url": url
        }
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload template: {str(e)}"
        )


@router.get("/template/download")
async def download_parts_template():
    """
    Download the parts template file from MinIO.
    """
    minio_client = get_minio_client()

    try:
        # Check if template exists
        if not minio_client.file_exists(PARTS_TEMPLATE_OBJECT_NAME):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Parts template not found. Please upload the template first."
            )

        # Download from MinIO
        content = minio_client.download_file(PARTS_TEMPLATE_OBJECT_NAME)

        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=PartsTemplate.docx"}
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to download template: {str(e)}"
        )


@router.get("/template/exists")
async def check_parts_template_exists():
    """
    Check if the parts template exists in MinIO.
    """
    minio_client = get_minio_client()

    try:
        exists = minio_client.file_exists(PARTS_TEMPLATE_OBJECT_NAME)
        return {"exists": exists}
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to check template: {str(e)}"
        )
